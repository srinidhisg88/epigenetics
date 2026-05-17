"""
Generates the full 60-page project report for the
Epilepsy Diagnostic Assistant project as a .docx file.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page setup: exact A4 (210 × 297 mm) ─────────────────────────────────────
section = doc.sections[0]
section.page_height   = Mm(297)   # exact A4 height
section.page_width    = Mm(210)   # exact A4 width
section.left_margin   = Mm(25)    # 25 mm left
section.right_margin  = Mm(25)    # 25 mm right
section.top_margin    = Mm(25)    # 25 mm top
section.bottom_margin = Mm(25)    # 25 mm bottom


# ─── Page border helper ───────────────────────────────────────────────────────
def add_page_border(section):
    """Add a box border around every page via sectPr XML."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '18')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

add_page_border(section)

# ─── Enforce A4 paper size in sectPr XML ─────────────────────────────────────
def set_a4_paper_size(section):
    """Set w:pgSz to A4 (11906 × 16838 twentieths-of-a-point = 210×297 mm)."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:pgSz')):
        sectPr.remove(existing)
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')   # 210 mm in twentieths-of-a-point
    pgSz.set(qn('w:h'), '16838')   # 297 mm in twentieths-of-a-point
    pgSz.set(qn('w:orient'), 'portrait')
    sectPr.insert(0, pgSz)

set_a4_paper_size(section)

# ─── Header (chapter title — top right) & Footer (page number — bottom centre) ─
def setup_header_footer(section):
    """Right-aligned header with document title; centered footer with PAGE field."""
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # ── Footer ────────────────────────────────────────────────────────────────
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run()
    frun.font.name = 'Times New Roman'
    frun.font.size = Pt(11)
    # PAGE field code
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    frun._r.append(fldChar1)
    frun._r.append(instrText)
    frun._r.append(fldChar2)

    # Decimal page numbers (1, 2, 3…) for chapters section (body sectPr)
    sp = section._sectPr
    for old in sp.findall(qn('w:pgNumType')):
        sp.remove(old)
    pnt = OxmlElement('w:pgNumType')
    pnt.set(qn('w:fmt'), 'decimal')
    pnt.set(qn('w:start'), '1')
    sp.append(pnt)

setup_header_footer(section)

# Capture footer rId so inline section breaks can reference the same footer part
_FOOTER_RID = None
for _ref in doc.sections[0]._sectPr.findall(qn('w:footerReference')):
    if _ref.get(qn('w:type')) == 'default':
        _FOOTER_RID = _ref.get(qn('r:id'))
        break

# ─── Section-break helpers (for Roman / Arabic page numbering) ────────────────
def _make_inline_sectPr(fmt, start, include_footer):
    """Return a sectPr OxmlElement with A4 size, margins, borders, pgNumType."""
    sectPr = OxmlElement('w:sectPr')

    # Optionally link to the shared footer part
    if include_footer and _FOOTER_RID:
        fRef = OxmlElement('w:footerReference')
        fRef.set(qn('w:type'), 'default')
        fRef.set(qn('r:id'), _FOOTER_RID)
        sectPr.append(fRef)

    # Next-page break type
    typ = OxmlElement('w:type')
    typ.set(qn('w:val'), 'nextPage')
    sectPr.append(typ)

    # A4 page size
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    pgSz.set(qn('w:orient'), 'portrait')
    sectPr.append(pgSz)

    # Margins — 25 mm all sides (1418 twips), 10 mm header/footer
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'),    '1418')
    pgMar.set(qn('w:right'),  '1418')
    pgMar.set(qn('w:bottom'), '1418')
    pgMar.set(qn('w:left'),   '1418')
    pgMar.set(qn('w:header'), '567')
    pgMar.set(qn('w:footer'), '567')
    sectPr.append(pgMar)

    # Page borders (same style as document border)
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '18')
        b.set(qn('w:space'), '24')
        b.set(qn('w:color'), '000000')
        pgBorders.append(b)
    sectPr.append(pgBorders)

    # Page number format
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

    return sectPr

def title_page_section_break():
    """After title page: Section 1 ends here.
    No footer → title page shows NO page number.
    lowerRoman,start=1 so abstract (Section 2) begins at i."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p._p.get_or_add_pPr().append(
        _make_inline_sectPr('lowerRoman', 1, include_footer=False)
    )

def precontent_section_break():
    """After List of Tables: Section 2 ends here.
    Footer with PAGE field, lowerRoman start=1 → abstract=i … LOT=v.
    Chapter 1 (Section 3 / body sectPr) starts at Arabic 1."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p._p.get_or_add_pPr().append(
        _make_inline_sectPr('lowerRoman', 1, include_footer=True)
    )


# ─── Style helpers ────────────────────────────────────────────────────────────
def set_font(run, size_pt=12, bold=False, italic=False, color=None):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_font(para, size_pt=12):
    """Apply Times New Roman to every run already in para."""
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size_pt)

def heading1(text):
    """Chapter/Appendix heading — 'Chapter N' italic above, bold title centred, horizontal rule below."""
    import re

    ch_match  = re.match(r'^CHAPTER\s+(\d+):\s*(.+)$',  text, re.IGNORECASE)
    app_match = re.match(r'^APPENDIX\s+([A-Z]):\s*(.+)$', text, re.IGNORECASE)

    if ch_match:
        label = f"Chapter {ch_match.group(1)}"
        title = ch_match.group(2).upper()
    elif app_match:
        label = f"Appendix {app_match.group(1)}"
        title = app_match.group(2).upper()
    else:
        label = None
        title = text.upper()

    # ── "Chapter N" / "Appendix A" label in italic ──────────────────────────
    if label:
        p_lbl = doc.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_lbl.paragraph_format.space_before = Pt(18)
        p_lbl.paragraph_format.space_after  = Pt(0)
        r_lbl = p_lbl.add_run(label)
        set_font(r_lbl, 14, bold=True)
        r_lbl.font.italic = True

    # ── Title line ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if label else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2) if label else Pt(18)
    p.paragraph_format.space_after  = Pt(14)
    run = p.add_run(title)
    set_font(run, 14, bold=True)

    # ── Horizontal rule only for chapters/appendices ─────────────────────────
    if label:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)

    return p

def heading2(text):
    """Section heading — 12 pt bold, left."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, 12, bold=True)
    return p

def body(text, justify=True):
    """Body paragraph — 12 pt, justified, 1.5 line spacing."""
    from docx.shared import Pt as _Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = _Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, 12)
    return p

def heading3(text):
    """Sub-section heading — 12 pt bold italic, left."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 12, bold=True, italic=True)
    return p

def add_image(path, caption, width=5.8):
    """Embed an image with a centered caption below it."""
    import os, io
    from PIL import Image as _PIL
    if not os.path.exists(path):
        return
    # Re-save with 150 DPI so python-docx renders at correct physical size
    img = _PIL.open(path)
    buf = io.BytesIO()
    img.save(buf, format='PNG', dpi=(150, 150))
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    set_font(r, 11, bold=True, italic=True)

def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.keep_with_next = False
    # Prevent paragraph from splitting across pages
    pPr = p._p.get_or_add_pPr()
    widowControl = OxmlElement('w:widowControl')
    widowControl.set(qn('w:val'), '0')
    pPr.append(widowControl)
    run = p.add_run(u"\u2022  " + text)
    set_font(run, 12)
    return p

def page_break():
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
run = p.add_run("A PROJECT REPORT")
set_font(run, 16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ON")
set_font(run, 14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
run = p.add_run("AI-POWERED EPILEPSY VARIANT DIAGNOSTIC ASSISTANT")
set_font(run, 16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Using Machine Learning, ACMG/AMP Variant Classification and Retrieval-Augmented Generation")
set_font(run, 12, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Submitted in partial fulfillment of the requirements for the award of the degree of")
set_font(run, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BACHELOR OF ENGINEERING")
set_font(run, 13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("in")
set_font(run, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Computer Science and Engineering")
set_font(run, 13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Submitted by")
set_font(run, 12, bold=True)

for name in ["Srinidhi S G", "Bharath H", "Sushant Kabbinada", "Aditya Patil"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    set_font(run, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
run = p.add_run("Under the Guidance of")
set_font(run, 12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dr. M P Pushpalatha")
set_font(run, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Associate Professor, Dept. of Computer Science & Engineering")
set_font(run, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING")
set_font(run, 13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("JSS Science and Technology University")
set_font(run, 13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mysuru – 570 006, Karnataka, India")
set_font(run, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Academic Year: 2025 – 2026")
set_font(run, 12, bold=True)

title_page_section_break()   # title page -> Section 1 (no footer shown)


# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
heading1("ABSTRACT")
body(
    "Epilepsy is one of the most prevalent neurological disorders worldwide, affecting over 50 million individuals. "
    "Accurate classification of genetic variants associated with epilepsy is essential for diagnosis, treatment planning, "
    "and genetic counseling. Manual curation of variants against established guidelines such as ACMG/AMP is time-intensive, "
    "inconsistent across laboratories, and requires deep domain expertise. This project presents an AI-powered Epilepsy "
    "Variant Diagnostic Assistant that integrates machine learning, automated ACMG/AMP classification, and "
    "Retrieval-Augmented Generation (RAG) to provide accurate, explainable variant pathogenicity predictions."
)
body(
    "The system employs an XGBoost classifier calibrated with isotonic regression over a 93-dimensional feature vector "
    "derived from genomic, clinical, and database features. The model achieves an accuracy of 89.89%, an AUC of 0.9446, "
    "and an F1 score of 0.8538 on held-out test data. An automated ACMG/AMP engine implements 14 evidence criteria using "
    "the Tavtigian et al. (2018) points-based scoring model, validated at 98.5% agreement with expert classifications "
    "across 200 real ClinVar variants. TreeSHAP is used to generate interpretable explanations and to supply PP3/BP4 "
    "evidence criteria. External data from ClinVar, gnomAD, PubMed, and PharmGKB is retrieved via a multi-source "
    "RAG pipeline and synthesized into clinical narrative reports using a large language model."
)
body(
    "A confidence-aware routing mechanism identifies uncertain predictions (probability 0.30–0.70) and triggers "
    "additional evidence gathering. A contradiction detection module identifies conflicts between evidence sources. "
    "The full system is deployed as a FastAPI backend with a React frontend, supporting real-time variant analysis "
    "for 27 epilepsy genes. Comprehensive testing across 200 test cases (unit, integration, regression, smoke, and "
    "performance) confirms system correctness, stability, and production readiness."
)
body(
    "Keywords: Epilepsy, Variant Classification, ACMG/AMP, XGBoost, SHAP, Retrieval-Augmented Generation, "
    "ClinVar, gnomAD, Pathogenicity Prediction, Clinical Decision Support."
)
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# ACKNOWLEDGEMENT
# ─────────────────────────────────────────────────────────────────────────────
heading1("ACKNOWLEDGEMENT")
body(
    "We would like to express our sincere gratitude to JSS Science and Technology University for providing us "
    "with the opportunity and resources to undertake this project, \u201cAI-Driven Epilepsy Genetic Variant "
    "Classification With ACMG Criteria and Retrieval-Augmented Generation.\u201d First and foremost, we extend "
    "our heartfelt appreciation to our guide, Dr. M P Pushpalatha, Professor, Department of Computer Science "
    "& Engineering, for her invaluable guidance, encouragement, and unwavering support throughout the course "
    "of this research. Her expertise and insights have been instrumental in shaping the direction of our work. "
    "We are also deeply grateful to Dr. Anil Kumar, Professor and Head of the Department, for his continuous "
    "support and for fostering an environment conducive to research and innovation."
)
body(
    "Our sincere thanks go to JSS Science and Technology University for providing us with access to essential "
    "resources, infrastructure, and academic support that enabled us to carry out this study effectively. We "
    "would also like to acknowledge the contributions of our fellow researchers, faculty members, and peers "
    "who provided valuable feedback and suggestions that helped refine our approach. Finally, we extend our "
    "deepest gratitude to our families and friends for their unwavering encouragement and support throughout "
    "this journey. Their belief in our work has been a constant source of motivation. This project has been "
    "a remarkable learning experience, and we are grateful to everyone who played a role in its successful "
    "completion."
)
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS  (tabular, bordered)
# ─────────────────────────────────────────────────────────────────────────────
heading1("TABLE OF CONTENTS")

# (entry text, page label, is_chapter_heading, indent_level)
toc_entries = [
    ("Abstract",                                              "i",       False, 0),
    ("Acknowledgement",                                       "ii",      False, 0),
    ("Table of Contents",                                     "iii",     False, 0),
    ("List of Figures",                                       "iv",      False, 0),
    ("List of Tables",                                        "v",       False, 0),
    ("Chapter 1: Introduction",                               "1",       True,  0),
    ("1.1 Overview of the Chapter",                          "1",       False, 1),
    ("1.2 Problem Statement",                                 "2",       False, 1),
    ("1.3 Aim",                                               "3",       False, 1),
    ("1.4 Objectives",                                        "3",       False, 1),
    ("1.5 Applications",                                      "4",       False, 1),
    ("1.6 Existing Solutions",                                "5",       False, 1),
    ("1.7 Proposed Solution",                                 "6",       False, 1),
    ("1.8 Gantt Chart",                                       "7",       False, 1),
    ("Chapter 2: Literature Review with Gap in the Literature","8",      True,  0),
    ("2.1 – 2.7  Literature Sections",                        "8–12",    False, 1),
    ("2.8 Limitations of the Present Work",                   "13",      False, 1),
    ("Chapter 3: System Requirements and Analysis",           "14",      True,  0),
    ("3.1 Hardware Requirements",                             "14",      False, 1),
    ("3.2 Software Requirements",                             "15",      False, 1),
    ("3.3 Functional Requirements",                           "16",      False, 1),
    ("3.4 Non-Functional Requirements",                       "17",      False, 1),
    ("Chapter 4: Tools and Technology Used",                  "19",      True,  0),
    ("Chapter 5: System Design",                              "26",      True,  0),
    ("5.1 Overview of System Workflow",                       "26",      False, 1),
    ("5.2 System Architecture Layers",                        "27",      False, 1),
    ("5.3 ACMG Criterion Points",                             "30",      False, 1),
    ("5.4 Classification Thresholds",                         "31",      False, 1),
    ("5.5 Flow Diagram",                                      "32",      False, 1),
    ("Chapter 6: System Implementation",                      "34",      True,  0),
    ("6.1 ML Pipeline Implementation",                        "34",      False, 1),
    ("6.2 ACMG Engine Implementation",                        "39",      False, 1),
    ("6.3 RAG Pipeline Implementation",                       "43",      False, 1),
    ("6.4 User Interface Implementation",                     "46",      False, 1),
    ("Chapter 7: Testing and Results",                        "48",      True,  0),
    ("7.1 Testing Methodology",                               "48",      False, 1),
    ("7.2 Test Fixtures",                                     "49",      False, 1),
    ("7.3 Smoke Tests",                                       "49",      False, 1),
    ("7.4 Unit Tests",                                        "50",      False, 1),
    ("7.5 Integration Tests",                                 "51",      False, 1),
    ("7.6 Regression Tests",                                  "52",      False, 1),
    ("7.7 Performance Tests",                                 "53",      False, 1),
    ("7.8 Overall Test Results Summary",                      "55",      False, 1),
    ("Chapter 8: Conclusion and Future Work",                 "56",      True,  0),
    ("8.1 Conclusion",                                        "56",      False, 1),
    ("8.2 Future Work",                                       "58",      False, 1),
    ("Appendix A: Project Team Details",                      "60",      True,  0),
    ("Appendix B: COs, POs and PSOs Mapping",                 "61",      True,  0),
    ("Appendix C: Publication Details",                       "64",      True,  0),
    ("References",                                            "65",      True,  0),
]

def add_outer_border(tbl, color='000000', sz='6'):
    """Add outer border only to a table via tblBorders XML."""
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    # insideH and insideV — none
    for side in ('insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

toc_tbl = doc.add_table(rows=len(toc_entries), cols=2)
toc_tbl.style = 'Normal Table'
toc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_outer_border(toc_tbl)
# Set column widths
from docx.oxml import OxmlElement as _OXE
toc_tbl.columns[0].width = Mm(130)
toc_tbl.columns[1].width = Mm(20)

for i, (entry, pg, is_chap, indent) in enumerate(toc_entries):
    row = toc_tbl.rows[i]
    # Left cell — chapter/section title
    lc = row.cells[0]
    lp = lc.paragraphs[0]
    lp.clear()
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after  = Pt(2)
    lp.paragraph_format.left_indent  = Inches(indent * 0.3)
    lr = lp.add_run(entry)
    set_font(lr, 12, bold=is_chap)
    # Right cell — page number centred
    rc = row.cells[1]
    rp = rc.paragraphs[0]
    rp.clear()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.paragraph_format.space_before = Pt(2)
    rp.paragraph_format.space_after  = Pt(2)
    rr = rp.add_run(pg)
    set_font(rr, 12, bold=is_chap)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# LIST OF FIGURES  (tabular — Figure No. | Figure Description | Page No.)
# ─────────────────────────────────────────────────────────────────────────────
heading1("LIST OF FIGURES")

figures = [
    ("5.1",  "System Data Pipeline Flowchart — ClinVar Download, VUS/UTR Filtering, Stratified Split and SMOTE+Tomek Resampling", "32"),
    ("5.2",  "Module Chain — Blue Boxes Execute on Every Request; Orange Boxes Activate for Pathogenic or Uncertain Predictions",  "33"),
    ("5.3",  "Ten-Layer System Architecture — L1 React Frontend to L10 PredictionResult JSON",                                    "34"),
    ("6.1",  "Class Distribution per Gene Before and After SMOTE Balancing",                                                      "36"),
    ("6.2",  "Feature Engineering Pipeline — Raw Input to 93-Dimensional Feature Vector",                                         "37"),
    ("6.3",  "Model Training Workflow — SMOTE+Tomek, XGBoost Fitting, Isotonic Calibration and Val/Test Scoring",                 "39"),
    ("6.4",  "ROC Curve (AUC = 0.9446) and Confusion Matrix on Held-Out Test Set",                                               "41"),
    ("6.5",  "SHAP Feature Importance — Top Predictive Features for Pathogenicity",                                               "42"),
    ("6.6",  "Variant Pathogenicity Analysis — Empty Input Form with Quick Example Buttons and No Prediction Yet",               "47"),
    ("6.7",  "SCN1A C→T Nonsense Variant — Filled Form with Pathogenic Result (99.5% Confidence) and ACMG Score +9",            "47"),
    ("6.8",  "ACMG/AMP Classification Panel — Evidence Score +9, PVS1 Very Strong and PP3 Supporting Criteria",                 "48"),
    ("6.9",  "Diagnostic Results Tab — Variant Interpretation, Contradictory Evidence and Why This Variant Matters Clinically",  "48"),
    ("6.10", "RAG Follow-up Chat — Contraindicated Medications and Gene-Specific LLM Clinical Reasoning",                        "49"),
]

lof_tbl = doc.add_table(rows=len(figures)+1, cols=3)
lof_tbl.style = 'Normal Table'
lof_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_outer_border(lof_tbl)
lof_tbl.columns[0].width = Mm(20)
lof_tbl.columns[1].width = Mm(115)
lof_tbl.columns[2].width = Mm(15)

for j, h in enumerate(["Figure No.", "Figure Description", "Page No."]):
    cell = lof_tbl.rows[0].cells[j]
    p = cell.paragraphs[0]; p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    set_font(r, 12, bold=True)

for i, (fig_no, desc, pg) in enumerate(figures):
    row = lof_tbl.rows[i+1]
    # Figure No.
    p0 = row.cells[0].paragraphs[0]; p0.clear()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(fig_no); set_font(r0, 12)
    # Description
    p1 = row.cells[1].paragraphs[0]; p1.clear()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(desc); set_font(r1, 12)
    # Page No.
    p2 = row.cells[2].paragraphs[0]; p2.clear()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(pg); set_font(r2, 12)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# LIST OF TABLES  (tabular — Table No. | Table Description | Page No.)
# ─────────────────────────────────────────────────────────────────────────────
heading1("LIST OF TABLES")

tables_list = [
    ("1.1",  "Project Gantt Chart",                              "7"),
    ("2.1",  "Summary of Related Works",                         "13"),
    ("3.1",  "Hardware Requirements",                            "14"),
    ("3.2",  "Software Requirements",                            "15"),
    ("3.3",  "Functional Requirements",                          "16"),
    ("3.4",  "Non-Functional Requirements",                      "17"),
    ("4.1",  "Summary of Tools and Technologies",                "25"),
    ("5.1",  "ACMG/AMP Criterion Points (Tavtigian et al. 2018)","30"),
    ("5.2",  "Five-Tier Classification Score Thresholds",        "31"),
    ("6.1",  "Model Comparison — Accuracy, AUC, F1",             "35"),
    ("6.2",  "Epilepsy Gene Panel — Disease Mechanism",          "40"),
    ("7.1",  "Test Fixtures Used Across All Suites",             "49"),
    ("7.2",  "Performance SLA Compliance",                       "53"),
    ("7.3",  "Overall Test Results Summary",                     "55"),
    ("A.1",  "Project Team Details",                             "62"),
    ("B.1",  "Mapping of CO, PO and PSO",                        "64"),
    ("C.1",  "Publication Details",                              "66"),
]

lot_tbl = doc.add_table(rows=len(tables_list)+1, cols=3)
lot_tbl.style = 'Normal Table'
lot_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_outer_border(lot_tbl)
lot_tbl.columns[0].width = Mm(20)
lot_tbl.columns[1].width = Mm(115)
lot_tbl.columns[2].width = Mm(15)

for j, h in enumerate(["Table No.", "Table Description", "Page No."]):
    cell = lot_tbl.rows[0].cells[j]
    p = cell.paragraphs[0]; p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    set_font(r, 12, bold=True)

for i, (tbl_no, desc, pg) in enumerate(tables_list):
    row = lot_tbl.rows[i+1]
    p0 = row.cells[0].paragraphs[0]; p0.clear()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(tbl_no); set_font(r0, 12)
    p1 = row.cells[1].paragraphs[0]; p1.clear()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(desc); set_font(r1, 12)
    p2 = row.cells[2].paragraphs[0]; p2.clear()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(pg); set_font(r2, 12)

precontent_section_break()   # LOT ends Section 2 (roman i-v); Chapter 1 starts Section 3 (arabic 1)


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 1: INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 1: INTRODUCTION")

heading2("1.1 Overview of the Chapter")
body(
    "Epilepsy is a chronic neurological disorder characterised by recurrent, unprovoked seizures resulting from "
    "abnormal electrical activity in the brain. It is one of the most common serious neurological conditions globally, "
    "affecting more than 50 million people across all ages and geographies. In approximately 30–40% of epilepsy cases, "
    "the condition has a genetic aetiology, with pathogenic variants in ion channel genes, synaptic transmission genes, "
    "and mTOR pathway regulators playing a central role."
)
body(
    "The identification and accurate classification of genetic variants is a critical step in the diagnostic workup "
    "of patients with suspected genetic epilepsy. The American College of Medical Genetics and Genomics (ACMG) and the "
    "Association for Molecular Pathology (AMP) published standardised guidelines in 2015 for the interpretation of "
    "sequence variants, using a five-tier classification system: Pathogenic, Likely Pathogenic, Variant of Uncertain "
    "Significance (VUS), Likely Benign, and Benign. These guidelines were later extended by Tavtigian et al. (2018) "
    "using a Bayesian points-based scoring framework to provide unambiguous numerical thresholds."
)
body(
    "Despite the availability of these guidelines, the manual application of ACMG/AMP criteria is a labour-intensive "
    "process that requires integration of data from multiple databases (ClinVar, gnomAD, PubMed), domain knowledge of "
    "gene-disease mechanisms, and experienced clinical judgement. In practice, this leads to inconsistency across "
    "laboratories and significant delays in variant classification, impacting patient care."
)
body(
    "This chapter introduces the Epilepsy Variant Diagnostic Assistant, an AI-powered clinical decision support system "
    "that automates variant pathogenicity prediction and ACMG/AMP classification for a panel of 27 epilepsy-associated "
    "genes. The system integrates a calibrated XGBoost classifier, an automated ACMG/AMP engine, SHAP-based "
    "explainability, and a multi-source Retrieval-Augmented Generation (RAG) pipeline to produce interpretable, "
    "evidence-backed clinical reports."
)

heading2("1.2 Problem Statement")
body(
    "The classification of genetic variants found in patients with epilepsy presents several significant challenges "
    "in clinical genetics practice:"
)
bullet("Volume and velocity: Next-generation sequencing panels generate hundreds of candidate variants per patient, each requiring individual review against the ACMG/AMP framework.")
bullet("Manual inconsistency: Studies have shown inter-laboratory discordance rates of 10–20% for variant classification, particularly for Variants of Uncertain Significance (VUS).")
bullet("Multi-database integration: Accurate classification requires simultaneous lookup across ClinVar, gnomAD, PubMed literature, and gene-specific databases — a process that is time-consuming and error-prone when performed manually.")
bullet("Lack of explainability: Existing computational tools (CADD, REVEL, PolyPhen-2) produce scores without mapping their predictions directly to ACMG/AMP criteria, limiting clinical utility.")
bullet("Gene-specific knowledge gaps: Different epilepsy genes have distinct disease mechanisms (loss-of-function vs. gain-of-function), which critically affect which ACMG criteria are applicable — knowledge that is not captured in generic variant classifiers.")
body(
    "These challenges collectively slow down the diagnostic process, increase the proportion of unresolved VUS cases, "
    "and potentially delay life-altering treatment decisions for patients and families. There is a clear clinical need "
    "for an automated, explainable, multi-evidence system specifically designed for epilepsy variant classification."
)

heading2("1.3 Aim")
body(
    "The aim of this project is to design, develop, and validate an AI-powered Epilepsy Variant Diagnostic Assistant "
    "that automates the classification of genetic variants in epilepsy-associated genes according to ACMG/AMP guidelines, "
    "integrates multi-source evidence retrieval, and produces interpretable clinical reports suitable for use by clinical "
    "geneticists and neurologists."
)

page_break()
heading2("1.4 Objectives")
body("The specific objectives of this project are:")
bullet("To develop a machine learning model (XGBoost with isotonic calibration) that classifies genetic variants in 27 epilepsy genes as Pathogenic or Benign with accuracy above 89%.")
bullet("To engineer a 93-dimensional feature vector capturing genomic, clinical, database-derived, and gene-specific features for each variant.")
bullet("To implement an automated ACMG/AMP classification engine applying 14 evidence criteria (PVS1, PS1–PS3, PM1–PM6, PP3, PP5, BA1, BS1–BS2, BP4, BP6–BP7) using the Tavtigian points-based scoring model.")
bullet("To integrate TreeSHAP for generating interpretable feature attributions and mapping them to PP3/BP4 ACMG criteria.")
bullet("To build a multi-source Retrieval-Augmented Generation (RAG) pipeline integrating ClinVar REST API, gnomAD GraphQL API, PubMed FAISS index, and PharmGKB.")
bullet("To implement a confidence-aware routing mechanism that identifies uncertain predictions and triggers additional evidence gathering.")
bullet("To deploy a contradiction detection module that identifies conflicts between evidence sources.")
bullet("To build a React-based frontend and FastAPI backend providing a clinical web interface.")
bullet("To validate the system through comprehensive testing across 200 test cases spanning unit, integration, regression, smoke, and performance tests.")

heading2("1.5 Applications")
body(
    "The Epilepsy Variant Diagnostic Assistant has direct applications across clinical, research, and educational domains:"
)
bullet("Clinical Genetics Laboratories: Automates the initial classification of variants identified in epilepsy gene panel sequencing, reducing manual workload and improving throughput.")
bullet("Neurogenomics Clinics: Provides real-time decision support to clinical geneticists reviewing variants from whole-exome or whole-genome sequencing in epilepsy patients.")
bullet("Genetic Counseling: Enables genetic counselors to quickly assess variant pathogenicity and generate patient-friendly summaries for communication with families.")
bullet("Research and Variant Curation: Supports large-scale variant curation projects by providing automated ACMG/AMP classifications with full evidence documentation.")
bullet("Medical Education: Provides a teaching tool for clinical genetics trainees to understand the application of ACMG/AMP criteria to real epilepsy variants.")
bullet("Pharmacogenomics: Integration with PharmGKB enables identification of variants with drug-response implications for anti-epileptic drug selection.")
bullet("Hospital Information Systems: Can be integrated into electronic health records as a decision-support plugin for neurology departments.")

page_break()
heading2("1.6 Existing Solutions")
body(
    "Several tools and databases have been developed to assist with variant classification, but each has significant "
    "limitations when applied to the specific domain of epilepsy genetics:"
)
body(
    "ClinVar is the primary public repository of variant-disease relationships maintained by NCBI. While invaluable, "
    "ClinVar is a database rather than a classifier — it reports existing submissions rather than predicting "
    "pathogenicity for novel variants. Furthermore, a large proportion of epilepsy variants in ClinVar are classified "
    "as VUS, providing limited diagnostic utility."
)
body(
    "CADD (Combined Annotation-Dependent Depletion) and REVEL provide pathogenicity scores for missense variants but "
    "do not produce ACMG-tier classifications, do not account for gene-specific disease mechanisms, and lack "
    "integration with population frequency or de novo status evidence."
)
body(
    "InterVar is an automated tool for ACMG/AMP variant interpretation but is a general-purpose classifier that "
    "lacks domain-specific gene mechanism knowledge for epilepsy genes and does not integrate a RAG narrative "
    "generation layer."
)
body(
    "Epi25 and the Genetic Epilepsy Team at Boston Children's Hospital have published large epilepsy-specific variant "
    "datasets, but these are research databases rather than clinical decision support tools."
)
body(
    "The primary gap across all existing solutions is the absence of an integrated system that combines ML-based "
    "pathogenicity prediction, automated ACMG/AMP classification with gene-mechanism awareness, SHAP explainability "
    "mapped to ACMG criteria, and multi-source evidence RAG generation — all within a single clinical web application."
)

heading2("1.7 Proposed Solution")
body(
    "The proposed Epilepsy Variant Diagnostic Assistant is a ten-layer AI pipeline that addresses each of the "
    "identified gaps in existing solutions:"
)
bullet("Layer 1 — Input Validation: Pydantic-enforced schema validation of variant input (gene, chromosome, alleles, consequence, origin, review status).")
bullet("Layer 2 — Feature Engineering: Transformation of raw variant attributes into a 93-dimensional numerical feature vector including gene one-hot encoding, consequence severity flags, origin encoding, review score, and chromosomal position features.")
bullet("Layer 3 — ML Classifier: XGBoost (n_estimators=300, max_depth=8, lr=0.05) with isotonic calibration producing well-calibrated pathogenic probability estimates.")
bullet("Layer 4 — SHAP Explainer: TreeSHAP computes per-feature attributions translated into clinical language and mapped to ACMG PP3/BP4 criteria.")
bullet("Layer 5 — ACMG Engine: Automated application of 14 ACMG/AMP criteria with Tavtigian scoring, producing a five-tier classification with full criterion justification.")
bullet("Layer 6 — Multi-Source Retriever: Parallel queries to ClinVar REST API, gnomAD GraphQL API, a PubMed FAISS vector index (363 vectors), and PharmGKB.")
bullet("Layer 7 — Confidence Resolver: Six-zone routing based on pathogenic probability; triggers additional evidence gathering for uncertain predictions (0.30–0.70).")
bullet("Layer 8 — Contradiction Detector: Identifies conflicts between ML prediction, ClinVar classification, gnomAD frequency, and variant consequence.")
bullet("Layer 9 — RAG Generator: Synthesizes all evidence into a structured clinical narrative using an LLM (Groq llama-3.3-70b-versatile).")
bullet("Layer 10 — Response Schema: Structured JSON response with prediction, probabilities, confidence, SHAP explanation, gnomAD data, ACMG classification, and RAG narrative.")

page_break()
heading2("1.8 Gantt Chart")
body("Table 1.1 presents the project timeline from October 2025 to May 2026.")

p = doc.add_paragraph()
run = p.add_run("Table 1.1: Project Gantt Chart")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _shade_cell(cell, hex_color):
    """Fill a table cell with a solid background color (hex without #, e.g. '1F3864')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

phases = [
    ("Literature Review & Problem Formulation",  [1,1,0,0,0,0,0,0]),
    ("Dataset Collection & Preprocessing",        [0,1,1,0,0,0,0,0]),
    ("Feature Engineering",                       [0,1,1,0,0,0,0,0]),
    ("ML Model Development & Training",           [0,0,1,1,0,0,0,0]),
    ("ACMG Engine Implementation",               [0,0,0,1,1,0,0,0]),
    ("SHAP Explainer Integration",               [0,0,0,0,1,0,0,0]),
    ("RAG Pipeline Development",                  [0,0,0,0,1,1,0,0]),
    ("Frontend Development",                      [0,0,0,0,0,1,1,0]),
    ("System Integration & Testing",              [0,0,0,0,0,0,1,1]),
    ("Report Writing & Finalization",             [0,0,0,0,0,0,0,1]),
]
# Row colours (alternating shades of blue-green for clarity)
ROW_COLORS = [
    "2E75B6","2E75B6","2E75B6","2E75B6","2E75B6",
    "2E75B6","2E75B6","2E75B6","2E75B6","2E75B6",
]
months = ["Oct 2025","Nov 2025","Dec 2025","Jan 2026","Feb 2026","Mar 2026","Apr 2026","May 2026"]

tbl = doc.add_table(rows=len(phases)+2, cols=len(months)+1)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# ── Set column widths ──────────────────────────────────────────────────────
from docx.oxml.ns import qn as _qn
tblPr = tbl._tbl.find(_qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl._tbl.insert(0, tblPr)
tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'), '9360')   # total table width in twips (~16.5cm)
tblW.set(qn('w:type'), 'dxa')
tblPr.append(tblW)

# Column widths: activity col = 3000 twips, each month = 795 twips
tblGrid = OxmlElement('w:tblGrid')
for w in [3000] + [795]*len(months):
    gc = OxmlElement('w:gridCol')
    gc.set(qn('w:w'), str(w))
    tblGrid.append(gc)
tbl._tbl.insert(1, tblGrid)

# ── Title row (merged across all cols) ────────────────────────────────────
title_row = tbl.rows[0]
title_cell = title_row.cells[0]
# Merge all cells in the title row
for k in range(1, len(months)+1):
    title_cell = title_cell.merge(title_row.cells[k])
title_cell.text = "Project Schedule — October 2025 to May 2026"
for para in title_cell.paragraphs:
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        set_font(run, 11, bold=True)
_shade_cell(title_cell, "1F3864")
for para in title_cell.paragraphs:
    for run in para.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

# ── Header row ────────────────────────────────────────────────────────────
hdr = tbl.rows[1]
hdr.cells[0].text = "Activity / Phase"
for j, m in enumerate(months):
    hdr.cells[j+1].text = m
for cell in hdr.cells:
    _shade_cell(cell, "2E4057")
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, 9, bold=True)
            run.font.color.rgb = RGBColor(255, 255, 255)

# ── Data rows ─────────────────────────────────────────────────────────────
for i, (phase, sched) in enumerate(phases):
    row = tbl.rows[i+2]
    # Alternating row background for activity column
    bg = "EBF3FB" if i % 2 == 0 else "D6E8F7"
    _shade_cell(row.cells[0], bg)
    row.cells[0].text = phase
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            set_font(run, 9)
    for j, active in enumerate(sched):
        cell = row.cells[j+1]
        if active:
            _shade_cell(cell, "2E75B6")   # solid blue bar
            # Add a non-breaking space so the cell height renders
            for para in cell.paragraphs:
                para.clear()
                r = para.add_run("\u00a0")
                r.font.size = Pt(11)
        else:
            _shade_cell(cell, "F2F2F2")   # light grey for inactive

# Prevent Gantt rows from splitting across pages
def _set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), '1')
    trPr.append(cantSplit)

for row in tbl.rows:
    _set_row_cant_split(row)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 2: LITERATURE REVIEW
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 2: LITERATURE REVIEW WITH GAP IN THE LITERATURE")
body(
    "This chapter reviews existing research in the areas of genetic variant classification, machine learning "
    "for genomics, ACMG/AMP guidelines automation, explainable AI in clinical genomics, and Retrieval-Augmented "
    "Generation for medical knowledge synthesis. The review identifies key contributions and highlights the gaps "
    "that motivate the proposed system."
)

heading2("2.1 Variant Classification and ACMG/AMP Guidelines")
body(
    "Richards et al. (2015) published the ACMG/AMP Standards and Guidelines for the Interpretation of Sequence "
    "Variants, which established the five-tier classification system (Pathogenic, Likely Pathogenic, VUS, Likely "
    "Benign, Benign) and defined 28 evidence criteria grouped into very strong, strong, moderate, and supporting "
    "categories. These guidelines represented a landmark standardization effort, but their manual application "
    "remains subjective and time-consuming, particularly for the integration of population frequency data and "
    "computational evidence."
)
body(
    "Tavtigian et al. (2018) extended the ACMG/AMP framework by proposing a Bayesian points-based scoring model "
    "that assigns numerical weights to each criterion (e.g., PVS1 = 8 points, PS = 4 points, PM = 2 points, "
    "PP = 1 point, and their benign equivalents). This scoring model eliminates ambiguity in combining evidence "
    "and enables unambiguous numerical thresholds (≥10 = Pathogenic, 6–9 = Likely Pathogenic, 1–5 = VUS, "
    "-1 to -6 = Likely Benign, ≤-7 or BA1 = Benign). Our system directly implements this points model."
)

heading2("2.2 Machine Learning for Variant Pathogenicity Prediction")
body(
    "Kircher et al. (2014) introduced CADD (Combined Annotation-Dependent Depletion), a framework that integrates "
    "over 60 genomic annotations into a single deleteriousness score using a support vector machine trained to "
    "distinguish simulated de novo variants from those that have passed evolutionary selection. CADD has become "
    "a widely used computational tool but produces a continuous score rather than an ACMG-tier classification "
    "and lacks gene-specific disease mechanism knowledge."
)
body(
    "Ioannidis et al. (2016) developed REVEL (Rare Exome Variant Ensemble Learner), an ensemble method that "
    "combines scores from thirteen component pathogenicity predictors (including MutPred, FATHMM, PolyPhen-2, "
    "SIFT, and VEST) into a single meta-score. REVEL demonstrated improved prediction accuracy over individual "
    "tools for rare missense variants, but is restricted to missense variants and does not handle loss-of-function "
    "or splice variants — a critical limitation for epilepsy genes where frameshift, nonsense, and splice variants "
    "are among the most pathogenic classes."
)
body(
    "Landrum et al. (2018) described ClinVar, NCBI's curated database of variant-disease relationships. ClinVar "
    "provides aggregated clinical classifications from submitting laboratories worldwide and is an essential "
    "evidence source for PP5 and BP6 ACMG criteria. However, a substantial fraction of epilepsy gene variants "
    "remain unclassified or classified as VUS in ClinVar, limiting its utility as a standalone classification tool."
)
body(
    "Chen and Guestrin (2016) introduced XGBoost, a gradient boosting framework that has demonstrated state-of-the-art "
    "performance on structured tabular data across diverse domains. Its ability to handle missing values, mixed feature "
    "types, and regularisation through L1/L2 penalties makes it particularly suitable for genomic variant "
    "classification where features are heterogeneous and training data is limited."
)

heading2("2.3 Explainability in Clinical AI")
body(
    "Lundberg and Lee (2017) introduced SHAP (SHapley Additive exPlanations), a unified framework for "
    "interpreting machine learning model predictions based on Shapley values from cooperative game theory. "
    "SHAP satisfies consistency and local accuracy properties that make it theoretically well-grounded. "
    "The TreeSHAP variant enables exact, polynomial-time computation of SHAP values for tree-based models "
    "such as XGBoost, making it practical for per-prediction explanations in clinical settings. In our "
    "system, TreeSHAP attributions are mapped to ACMG PP3/BP4 criteria when the top contributor meets a "
    "15% contribution threshold in the pathogenic or benign direction."
)
body(
    "Holzinger et al. (2017) argued that explainability is essential for AI adoption in clinical settings, "
    "as clinicians require not just accurate predictions but justifiable, auditable reasoning to support "
    "clinical decisions. Our SHAP integration directly addresses this requirement by providing "
    "clinician-readable feature attributions alongside ACMG criterion justifications."
)

page_break()
heading2("2.4 Retrieval-Augmented Generation")
body(
    "Lewis et al. (2020) introduced Retrieval-Augmented Generation (RAG), a framework that augments a "
    "parametric language model with a non-parametric retrieval component to condition generation on "
    "retrieved documents. RAG significantly improves factual accuracy and enables language models to "
    "cite up-to-date evidence that was not present in their training data. In the medical domain, "
    "grounding LLM outputs in retrieved biomedical literature is essential to prevent hallucination "
    "of clinical facts."
)
body(
    "Xiong et al. (2023) demonstrated the application of RAG to biomedical question answering using "
    "PubMed-indexed literature, showing that domain-specific retrieval from curated biomedical corpora "
    "substantially outperforms closed-book generation. Our system constructs a FAISS vector index over "
    "363 epilepsy-specific PubMed abstracts using the PubMedBERT embedding model "
    "(pritamdeka/S-PubMedBert-MS-MARCO), providing targeted retrieval for each variant's gene."
)
body(
    "Johnson et al. (2023) reviewed the use of large language models in clinical genomics, noting that "
    "LLMs trained on general text corpora have limited ability to accurately interpret variant pathogenicity "
    "without grounding in structured clinical databases. Our multi-source retriever addresses this by "
    "integrating ClinVar, gnomAD, PubMed, and PharmGKB data before generation, ensuring the LLM narrative "
    "is anchored in verified clinical evidence."
)

heading2("2.5 Epilepsy-Specific Genomics")
body(
    "Scheffer et al. (2017) published the ILAE (International League Against Epilepsy) classification of "
    "epilepsies and epileptic seizures, establishing a framework that maps specific gene variants to "
    "epilepsy syndromes. Key genes include SCN1A (Dravet syndrome), KCNQ2 (neonatal epilepsy), TSC1/TSC2 "
    "(tuberous sclerosis complex), STXBP1 (Ohtahara syndrome), and SLC2A1 (GLUT1 deficiency syndrome). "
    "The distinction between loss-of-function (LOF) and gain-of-function (GOF) mechanisms — a critical "
    "determinant of PVS1 applicability — is well established for these genes and is encoded directly "
    "in our GENE_MECHANISM map."
)
body(
    "Nolan et al. (2019) conducted a systematic analysis of variant classification discordance across "
    "epilepsy genetics laboratories, finding that VUS classification rates exceeded 30% for many "
    "commonly sequenced genes. The high VUS burden in epilepsy genetics underscores the need for "
    "automated tools that can leverage multiple evidence sources to resolve uncertain classifications."
)

heading2("2.6 Population Frequency Databases")
body(
    "Karczewski et al. (2020) described gnomAD v3, the Genome Aggregation Database containing "
    "population-level allele frequencies from over 140,000 individuals. gnomAD is the primary "
    "source for ACMG PM2 (absent from population databases), BS1 (allele frequency greater than "
    "expected for disease), and BA1 (stand-alone benign, AF > 5%) criteria. Our gnomAD fetcher "
    "queries the gnomAD GraphQL API and caches results in SQLite with a 30-day TTL."
)

heading2("2.7 Related Work Summary")
body("Table 2.1 summarises the key related works and their limitations.")

p = doc.add_paragraph()
run = p.add_run("Table 2.1: Summary of Related Works")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rw_data = [
    ("Richards et al., 2015", "ACMG/AMP guidelines", "Manual, subjective, no ML integration"),
    ("Tavtigian et al., 2018", "Bayesian points-based ACMG scoring", "No automated implementation provided"),
    ("Kircher et al., 2014 (CADD)", "SVM-based deleteriousness score", "No ACMG tier, not epilepsy-specific"),
    ("Ioannidis et al., 2016 (REVEL)", "Ensemble of 13 predictors", "Missense only, no LOF/splice"),
    ("Chen & Guestrin, 2016 (XGBoost)", "Gradient boosting framework", "General purpose, no clinical mapping"),
    ("Lundberg & Lee, 2017 (SHAP)", "Shapley-based explainability", "Tool only, no clinical integration"),
    ("Lewis et al., 2020 (RAG)", "Retrieval-augmented generation", "General purpose, no genomics focus"),
    ("Karczewski et al., 2020 (gnomAD)", "Population frequency database", "Data source only, no classifier"),
]
rw_tbl = doc.add_table(rows=len(rw_data)+1, cols=3)
rw_tbl.style = 'Table Grid'
rw_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ["Reference", "Contribution", "Limitation"]
for j, h in enumerate(hdrs):
    rw_tbl.rows[0].cells[j].text = h
    for para in rw_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 10, bold=True)
for i, (ref, cont, lim) in enumerate(rw_data):
    row = rw_tbl.rows[i+1]
    for j, txt in enumerate([ref, cont, lim]):
        row.cells[j].text = txt
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 10)

page_break()
heading2("2.8 Limitations of the Present Work")
body(
    "Based on the literature review, the following gaps and limitations in existing work are identified, "
    "which the proposed system addresses:"
)
bullet("No existing tool provides ACMG/AMP classification specifically tailored to the gene-disease mechanism knowledge of epilepsy genes (LOF vs. GOF distinction for PVS1 applicability).")
bullet("Existing ML tools (CADD, REVEL) do not produce ACMG-tier classifications or map their predictions to evidence criteria, limiting direct clinical utility.")
bullet("No tool integrates ML prediction, SHAP explainability, multi-source database retrieval, ACMG classification, and RAG narrative generation within a single deployable clinical application.")
bullet("Current RAG systems for genomics do not integrate real-time ClinVar and gnomAD data into their retrieval context, risking outdated information in generated narratives.")
bullet("Confidence-aware routing and contradiction detection between evidence sources are absent from existing variant classification tools.")
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 3: SYSTEM REQUIREMENTS AND ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 3: SYSTEM REQUIREMENTS AND ANALYSIS")
body(
    "This chapter describes the hardware, software, functional, and non-functional requirements of the "
    "Epilepsy Variant Diagnostic Assistant. Requirements were elicited through analysis of clinical "
    "workflows, review of existing variant classification tools, and iterative consultation with the "
    "project guide."
)

heading2("3.1 Hardware Requirements")
body("Table 3.1 lists the minimum and recommended hardware specifications for deployment.")
p = doc.add_paragraph()
run = p.add_run("Table 3.1: Hardware Requirements")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
hw_data = [
    ("Processor", "Intel Core i5 (8th Gen) or equivalent", "Intel Core i7 / Apple M-series"),
    ("RAM", "8 GB", "16 GB or above"),
    ("Storage", "10 GB free space", "50 GB SSD"),
    ("GPU", "Not required (CPU inference)", "CUDA-capable GPU (for model retraining)"),
    ("Network", "Broadband (for ClinVar/gnomAD API)", "Stable high-speed internet"),
    ("Display", "1280 × 720 resolution", "1920 × 1080 or above"),
]
hw_tbl = doc.add_table(rows=len(hw_data)+1, cols=3)
hw_tbl.style = 'Table Grid'
hw_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Component", "Minimum", "Recommended"]):
    hw_tbl.rows[0].cells[j].text = h
    for para in hw_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, row_data in enumerate(hw_data):
    for j, txt in enumerate(row_data):
        hw_tbl.rows[i+1].cells[j].text = txt
        for para in hw_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 11)

heading2("3.2 Software Requirements")
body("Table 3.2 lists the software dependencies of the system.")
p = doc.add_paragraph()
run = p.add_run("Table 3.2: Software Requirements")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sw_data = [
    ("Operating System", "macOS 12+ / Ubuntu 20.04+ / Windows 10+", "—"),
    ("Python", "3.10 or above", "3.11.7"),
    ("Node.js", "18+", "20 LTS"),
    ("FastAPI", "0.104+", "0.115"),
    ("XGBoost", "1.7+", "2.0"),
    ("scikit-learn", "1.3+", "1.4"),
    ("SHAP", "0.43+", "0.44"),
    ("FAISS-CPU", "1.7+", "1.8"),
    ("sentence-transformers", "2.2+", "2.7"),
    ("React", "18+", "18.2"),
    ("Groq API", "Cloud service", "llama-3.3-70b-versatile"),
    ("SQLite", "3.36+", "For caching"),
    ("Git", "2.30+", "For version control"),
]
sw_tbl = doc.add_table(rows=len(sw_data)+1, cols=3)
sw_tbl.style = 'Table Grid'
sw_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Component", "Requirement", "Version Used"]):
    sw_tbl.rows[0].cells[j].text = h
    for para in sw_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, row_data in enumerate(sw_data):
    for j, txt in enumerate(row_data):
        sw_tbl.rows[i+1].cells[j].text = txt
        for para in sw_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 11)

page_break()
heading2("3.3 Functional Requirements")
body("Table 3.3 presents the functional requirements of the system.")
p = doc.add_paragraph()
run = p.add_run("Table 3.3: Functional Requirements")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr_data = [
    ("FR-01", "Variant Input", "The system shall accept variant input including gene symbol, chromosome, reference allele, alternate allele, molecular consequence, variant type, review status, and origin."),
    ("FR-02", "ML Prediction", "The system shall classify each variant as Pathogenic or Benign with a calibrated pathogenic probability."),
    ("FR-03", "SHAP Explanation", "The system shall generate SHAP-based feature attributions for each prediction."),
    ("FR-04", "ACMG Classification", "The system shall evaluate 14 ACMG/AMP criteria and produce a five-tier classification with Tavtigian scoring."),
    ("FR-05", "ClinVar Retrieval", "The system shall retrieve variant classifications from ClinVar and integrate them into PP5/BP6 criteria."),
    ("FR-06", "gnomAD Retrieval", "The system shall retrieve population allele frequency data and apply PM2/BA1/BS1 criteria."),
    ("FR-07", "Literature Retrieval", "The system shall retrieve relevant PubMed abstracts via FAISS semantic search."),
    ("FR-08", "RAG Narrative", "The system shall generate a structured clinical narrative report using an LLM conditioned on retrieved evidence."),
    ("FR-09", "Confidence Routing", "The system shall identify uncertain predictions and trigger additional evidence resolution."),
    ("FR-10", "Contradiction Detection", "The system shall identify and report contradictions between evidence sources."),
    ("FR-11", "Chat Interface", "The system shall support a conversational interface for follow-up clinical queries."),
    ("FR-12", "Gene/Consequence Lists", "The system shall expose the supported gene panel and consequence types via API."),
]
fr_tbl = doc.add_table(rows=len(fr_data)+1, cols=3)
fr_tbl.style = 'Table Grid'
fr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["ID", "Requirement", "Description"]):
    fr_tbl.rows[0].cells[j].text = h
    for para in fr_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 10, bold=True)
for i, row_data in enumerate(fr_data):
    for j, txt in enumerate(row_data):
        fr_tbl.rows[i+1].cells[j].text = txt
        for para in fr_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 10)

heading2("3.4 Non-Functional Requirements")
p = doc.add_paragraph()
run = p.add_run("Table 3.4: Non-Functional Requirements")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
nfr_data = [
    ("NFR-01", "Performance", "ML inference latency < 2 seconds per variant; full RAG pipeline < 30 seconds."),
    ("NFR-02", "Accuracy", "ML model accuracy ≥ 88%; ACMG classification agreement ≥ 95% on validated variants."),
    ("NFR-03", "Reliability", "System uptime ≥ 99%; zero crashes on malformed input."),
    ("NFR-04", "Scalability", "API must handle ≥ 10 concurrent users without degradation."),
    ("NFR-05", "Explainability", "Every prediction must be accompanied by SHAP feature attributions and ACMG criterion justifications."),
    ("NFR-06", "Security", "Invalid inputs must return HTTP 422 (not 500); no stack traces in error responses."),
    ("NFR-07", "Maintainability", "Codebase must be modular, with each pipeline layer in a separate file."),
    ("NFR-08", "Portability", "System must be deployable on macOS, Linux, and Windows via pip install."),
    ("NFR-09", "Usability", "Frontend must be accessible to non-specialist clinical staff within 10 minutes of instruction."),
    ("NFR-10", "Data Currency", "ClinVar and gnomAD data must be cached with a 30-day TTL to balance freshness and performance."),
]
nfr_tbl = doc.add_table(rows=len(nfr_data)+1, cols=3)
nfr_tbl.style = 'Table Grid'
nfr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["ID", "Category", "Requirement"]):
    nfr_tbl.rows[0].cells[j].text = h
    for para in nfr_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 10, bold=True)
for i, row_data in enumerate(nfr_data):
    for j, txt in enumerate(row_data):
        nfr_tbl.rows[i+1].cells[j].text = txt
        for para in nfr_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 10)
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 4: TOOLS AND TECHNOLOGY USED
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 4: TOOLS AND TECHNOLOGY USED")
body(
    "This chapter describes the programming languages, frameworks, libraries, APIs, and development tools "
    "used in the implementation of the Epilepsy Variant Diagnostic Assistant."
)

heading2("4.1 Python 3.11")
body(
    "Python 3.11 was selected as the primary backend programming language due to its extensive ecosystem "
    "of scientific computing, machine learning, and web development libraries. Python's dynamic typing "
    "and interpreted nature accelerated prototyping, while its mature packaging ecosystem (pip, conda) "
    "simplified dependency management. Key Python features used include dataclasses, type hints, "
    "asyncio for concurrent API calls, and context managers for database connections."
)

heading2("4.2 XGBoost")
body(
    "XGBoost (eXtreme Gradient Boosting) is a gradient-boosting framework optimised for speed and "
    "performance on tabular data. It implements regularised gradient boosting with L1 and L2 "
    "regularisation, built-in handling of missing values, and parallel tree construction. "
    "The deployed model uses: n_estimators=300, max_depth=8, learning_rate=0.05, subsample=0.8, "
    "colsample_bytree=0.8, min_child_weight=3, gamma=0.1, reg_alpha=0.1, reg_lambda=1.0. "
    "XGBoost outperformed all other candidate models on the epilepsy variant dataset due to its "
    "ability to model complex non-linear interactions between genomic features."
)

heading2("4.3 scikit-learn and Isotonic Calibration")
body(
    "scikit-learn provides the CalibratedClassifierCV wrapper used to apply isotonic regression "
    "calibration to the XGBoost model. Calibration ensures that predicted probabilities are "
    "statistically reliable — that is, variants predicted with 80% pathogenic probability are "
    "truly pathogenic approximately 80% of the time. Isotonic calibration was selected over "
    "Platt scaling because the XGBoost probability outputs were found to be non-linearly miscalibrated, "
    "and isotonic regression is a non-parametric method that does not assume a specific calibration curve shape."
)

heading2("4.4 SHAP (SHapley Additive exPlanations)")
body(
    "SHAP is used for generating per-prediction feature attributions via the TreeSHAP algorithm. "
    "SHAP values satisfy the properties of local accuracy (attributions sum to the prediction), "
    "missingness (zero attribution for absent features), and consistency (monotonicity with model behavior). "
    "In the system, SHAP attributions are translated into clinical language using a feature description "
    "dictionary and mapped to ACMG criteria: if the top contributor drives pathogenic prediction with "
    "≥15% contribution, PP3 is activated; if it drives benign prediction, BP4 is activated."
)

heading2("4.5 FastAPI")
body(
    "FastAPI is a modern, high-performance Python web framework for building APIs with automatic "
    "data validation via Pydantic, automatic OpenAPI documentation generation, and asynchronous "
    "request handling. The backend exposes eight endpoints: /health, /predict_variant, /analyze_variant, "
    "/explain_variant, /chat, /literature/{gene}, /genes, and /consequences. FastAPI's Pydantic "
    "integration provides Layer 1 input validation, rejecting malformed requests with HTTP 422 "
    "before they reach the ML pipeline."
)

page_break()
heading2("4.6 FAISS (Facebook AI Similarity Search)")
body(
    "FAISS is a library for efficient similarity search in high-dimensional vector spaces. "
    "In this project, a FAISS flat L2 index stores 363 PubMed abstract embeddings (768-dimensional "
    "PubMedBERT vectors). During variant analysis, the gene name is used as a query to retrieve "
    "the top-5 most semantically similar abstracts, which are included in the LLM context window "
    "for RAG generation. FAISS enables sub-millisecond retrieval over the full index."
)

heading2("4.7 Sentence Transformers (PubMedBERT)")
body(
    "The sentence-transformers library is used to generate dense vector embeddings for literature "
    "retrieval. The embedding model pritamdeka/S-PubMedBert-MS-MARCO is fine-tuned on PubMed "
    "text using the MS-MARCO passage ranking dataset, producing embeddings that capture "
    "biomedical semantic similarity. This model was selected over general-purpose embedding "
    "models (e.g., all-MiniLM-L6-v2) due to its superior performance on biomedical text retrieval tasks."
)

heading2("4.8 ClinVar REST API")
body(
    "The NCBI ClinVar REST API (eutils.ncbi.nlm.nih.gov) is queried to retrieve clinical "
    "significance classifications for variants in the same gene. The fetcher implements "
    "SQLite-based caching with a 30-day TTL, exponential backoff for rate limit handling, "
    "and review star extraction from structured ClinVar summary XML. ClinVar data drives the "
    "PP5 (pathogenic ClinVar classification) and BP6 (benign ClinVar classification) ACMG criteria."
)

heading2("4.9 gnomAD GraphQL API")
body(
    "The gnomAD GraphQL API (gnomad.broadinstitute.org/api) provides population allele frequency "
    "data. GraphQL queries are used to retrieve allele count, allele number, allele frequency, "
    "and homozygote count for specific variants. The gnomAD fetcher classifies variants as absent "
    "(0 alleles), ultra-rare (AF < 1×10⁻⁵), rare (AF < 0.001), or common (AF ≥ 0.01), "
    "mapping these classifications to ACMG criteria PM2, BA1, BS1, and BS2."
)

heading2("4.10 Groq API (LLM)")
body(
    "The Groq API provides inference for the llama-3.3-70b-versatile large language model, "
    "used to generate the RAG clinical narrative. The LLM receives a structured prompt containing "
    "the variant details, ACMG classification, SHAP attributions, ClinVar data, gnomAD data, "
    "retrieved PubMed abstracts, and PharmGKB information. The system prompt restricts the LLM "
    "to epilepsy genetics scope, preventing hallucination of out-of-domain content."
)

heading2("4.11 React and TypeScript (Frontend)")
body(
    "The frontend is built with React 18 and TypeScript, providing a type-safe, component-based "
    "user interface. Key components include VariantForm (input), ResultCard (ML prediction display), "
    "ACMGPanel (ACMG classification), SHAPPanel (feature attribution), and ChatInterface (conversational Q&A). "
    "Tailwind CSS is used for styling. The frontend communicates with the FastAPI backend via "
    "RESTful HTTP requests using the Fetch API."
)

heading2("4.12 PharmGKB API")
body(
    "PharmGKB is queried for pharmacogenomics annotations — specifically, gene-drug relationships "
    "relevant to anti-epileptic drug selection. This data is included in the RAG context when "
    "available, enabling the system to flag variants with drug-response implications alongside "
    "their pathogenicity classification."
)

p = doc.add_paragraph()
run = p.add_run("Table 4.1: Summary of Tools and Technologies")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tool_data = [
    ("Python 3.11", "Backend programming language"),
    ("XGBoost 2.0", "ML classifier"),
    ("scikit-learn 1.4", "Calibration, preprocessing"),
    ("SHAP 0.44", "Explainability, ACMG PP3/BP4"),
    ("FastAPI 0.115", "REST API backend"),
    ("Pydantic", "Input validation (Layer 1)"),
    ("FAISS-CPU 1.8", "Vector similarity search"),
    ("sentence-transformers", "PubMedBERT embeddings"),
    ("ClinVar REST API", "Clinical significance data"),
    ("gnomAD GraphQL API", "Population frequency data"),
    ("Groq API (llama-3.3-70b)", "LLM for RAG generation"),
    ("React 18 + TypeScript", "Frontend UI"),
    ("Tailwind CSS", "Frontend styling"),
    ("PharmGKB API", "Pharmacogenomics data"),
    ("SQLite", "Local caching (ClinVar, gnomAD)"),
    ("joblib", "Model serialization (.pkl)"),
    ("pandas + numpy", "Feature engineering"),
    ("pytest 9.0", "Testing framework"),
]
tool_tbl = doc.add_table(rows=len(tool_data)+1, cols=2)
tool_tbl.style = 'Table Grid'
tool_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Tool / Technology", "Purpose"]):
    tool_tbl.rows[0].cells[j].text = h
    for para in tool_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, (tool, purpose) in enumerate(tool_data):
    tool_tbl.rows[i+1].cells[0].text = tool
    tool_tbl.rows[i+1].cells[1].text = purpose
    for cell in tool_tbl.rows[i+1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, 11)
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 5: SYSTEM DESIGN
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 5: SYSTEM DESIGN")

heading2("5.1 Overview of System Workflow")
body(
    "The Epilepsy Variant Diagnostic Assistant follows a two-path workflow depending on the ML model's "
    "confidence level. The fast path (benign-confident) returns an ML prediction and ACMG classification "
    "without triggering the full RAG pipeline. The full path (pathogenic or uncertain) activates multi-source "
    "retrieval, confidence resolution, contradiction detection, and LLM narrative generation."
)
body(
    "Fast Path (benign probability > 70%): The variant input is validated, features are engineered, the "
    "XGBoost model produces a high-confidence benign prediction, SHAP attributions are computed, the ACMG "
    "engine classifies using locally available evidence, and a structured response is returned. "
    "This path completes in under 2 seconds."
)
body(
    "Full RAG Path (pathogenic or uncertain): All fast-path steps are completed, then the confidence "
    "resolver evaluates the pathogenic probability zone. If pathogenic (> 70%) or uncertain (30–70%), "
    "the multi-source retriever is invoked in parallel: ClinVar is queried for existing classifications, "
    "gnomAD is queried for population frequency, the FAISS index is searched for relevant PubMed abstracts, "
    "and PharmGKB is queried for drug-gene relationships. The retrieved evidence is integrated into the "
    "ACMG classification, the contradiction detector analyses source conflicts, and the RAG generator "
    "produces a comprehensive clinical narrative. This path typically completes in 230–490 milliseconds "
    "for local components, with additional LLM generation time."
)

page_break()
heading2("5.2 System Architecture Layers")
body(
    "The system is structured as ten discrete, independently testable layers. This modular architecture "
    "ensures separation of concerns, facilitates unit testing, and enables individual layer replacement "
    "without affecting the rest of the pipeline."
)

layers = [
    ("Layer 1 — Input Validation", "Pydantic BaseModel (VariantInput) enforces required fields: gene, chromosome, reference_allele, alternate_allele, consequence, with optional fields variant_type, review_status, origin. Invalid requests return HTTP 422 with field-level error messages before any processing occurs."),
    ("Layer 2 — Feature Engineering", "The raw variant input is transformed into a 93-dimensional numerical feature vector. Features include: 27 gene one-hot encoded flags, 9 consequence severity binary flags (is_frameshift, is_nonsense, is_missense, is_splice, is_synonymous, is_inframe, is_start_loss, is_stop_loss, severe_consequence_count), 3 variant type features (is_snp, is_transition, is_transversion), 5 review features (review_score, has_expert_review, has_multiple_submitters, has_criteria_provided, num_submitters), 3 position features, 2 origin features (is_germline, is_de_novo), and 9 gene category features (is_sodium_channel, is_ion_channel, is_gaba_receptor, is_tsc_complex, is_x_linked, is_rett_related, is_mtor_pathway, is_synaptic, is_glut1)."),
    ("Layer 3 — ML Classifier", "The calibrated XGBoost model (CalibratedClassifierCV with isotonic method) receives the 93-dimensional vector and returns pathogenic_probability and benign_probability. The model uses n_estimators=300, max_depth=8, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8, min_child_weight=3, gamma=0.1, reg_alpha=0.1, reg_lambda=1.0. Isotonic calibration corrects the raw XGBoost probability outputs to be statistically well-calibrated."),
    ("Layer 4 — SHAP Explainer", "TreeSHAP computes exact SHAP values for the XGBoost model. The top-5 contributing features are extracted, their contributions are expressed as percentages of the total absolute attribution, and each feature is mapped to a clinical description using the FEATURE_DESCRIPTIONS dictionary. The SHAP explainer is loaded as a singleton at startup to avoid repeated initialization overhead."),
    ("Layer 5 — ACMG Engine", "The ACMGClassifier evaluates 14 criteria: PVS1 (LOF in LOF gene), PS1 (same AA change — manual review required), PS2 (confirmed de novo), PS3 (functional studies — not auto-evaluated), PM1 (hotspot domain — SHAP-proxy), PM2 (gnomAD absent/ultra-rare), PM4 (in-frame indel), PM6 (assumed de novo), PP3 (SHAP pathogenic), PP5 (ClinVar pathogenic), BA1 (AF > 5%), BS1 (AF 1–5%), BS2 (healthy individuals), BP4 (SHAP benign), BP6 (ClinVar benign), BP7 (synonymous). Tavtigian scoring maps these to a five-tier classification."),
    ("Layer 6 — Multi-Source Retriever", "The MultiSourceRetriever orchestrates parallel queries to ClinVar (eutils REST API, SQLite-cached 30-day TTL), gnomAD (GraphQL API, SQLite-cached 30-day TTL), PubMed FAISS index (top-5 abstracts by gene query), and PharmGKB (gene-drug relationships). Results are aggregated and passed to the ACMG engine and RAG generator."),
    ("Layer 7 — Confidence Resolver", "The ConfidenceResolver classifies the pathogenic probability into six zones: very_high (≥0.9), high (0.7–0.9), moderate_pathogenic (0.5–0.7), moderate_benign (0.3–0.5), high_benign (0.1–0.3), very_high_benign (<0.1). Uncertain predictions (0.3–0.7) trigger additional evidence context building, with a weighted vote from gnomAD, ClinVar, consequence, and SHAP to suggest a classification."),
    ("Layer 8 — Contradiction Detector", "The ContradictionDetector identifies four classes of contradiction: ML vs. ClinVar (opposite classifications), ML vs. gnomAD (common variant predicted pathogenic), ML vs. Consequence (benign predicted for LOF consequence), and ClinVar Internal Conflict (conflicting submitter interpretations). Contradictions are classified by severity (low, medium, high) and included in the LLM context."),
    ("Layer 9 — RAG Generator", "The RAGGenerator constructs a structured prompt containing all evidence contexts (ACMG classification, SHAP attribution, gnomAD data, ClinVar data, contradiction analysis, uncertainty analysis, PubMed abstracts, PharmGKB data) and queries the Groq LLM API. The system prompt restricts responses to epilepsy genetics scope. The generator is wrapped with error handling to return a graceful error message if the LLM API is unavailable."),
    ("Layer 10 — Response Schema", "The FullAnalysisResponse Pydantic model enforces the response structure: prediction, confidence, pathogenic_probability, benign_probability, variant_info, rag_response, sources, is_pathogenic, shap_explanation, gnomad_data, contradictions, uncertainty_analysis, acmg_classification."),
]
for name, desc in layers:
    heading2(name)
    body(desc)

page_break()
heading2("5.3 ACMG Criterion Points — Tavtigian Model")
p = doc.add_paragraph()
run = p.add_run("Table 5.1: ACMG/AMP Criterion Points (Tavtigian et al. 2018)")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
acmg_data = [
    ("PVS1","Very Strong Pathogenic","+8","LOF variant in LOF gene"),
    ("PS1","Strong Pathogenic","+4","Same amino acid change as known pathogenic"),
    ("PS2","Strong Pathogenic","+4","Confirmed de novo variant"),
    ("PS3","Strong Pathogenic","+4","Functional studies confirm damaging effect"),
    ("PM1","Moderate Pathogenic","+2","Located in mutational hotspot/critical domain"),
    ("PM2","Moderate Pathogenic","+2","Absent from gnomAD population database"),
    ("PM4","Moderate Pathogenic","+2","In-frame indel or stop-loss"),
    ("PM6","Moderate Pathogenic","+2","Assumed de novo (unconfirmed)"),
    ("PP3","Supporting Pathogenic","+1","ML model predicts damaging (SHAP)"),
    ("PP5","Supporting Pathogenic","+1","ClinVar: Pathogenic classification"),
    ("BA1","Stand-alone Benign","−8","Allele frequency > 5% in gnomAD"),
    ("BS1","Strong Benign","−4","Allele frequency > 1% in gnomAD"),
    ("BS2","Strong Benign","−4","Observed in healthy individuals"),
    ("BP4","Supporting Benign","−1","ML model predicts tolerated (SHAP)"),
    ("BP6","Supporting Benign","−1","ClinVar: Benign classification"),
    ("BP7","Supporting Benign","−1","Synonymous variant, no splice impact"),
]
acmg_tbl = doc.add_table(rows=len(acmg_data)+1, cols=4)
acmg_tbl.style = 'Table Grid'
acmg_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Code","Category","Points","Description"]):
    acmg_tbl.rows[0].cells[j].text = h
    for para in acmg_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 10, bold=True)
for i, row_data in enumerate(acmg_data):
    for j, txt in enumerate(row_data):
        acmg_tbl.rows[i+1].cells[j].text = txt
        for para in acmg_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 10)

heading2("5.4 Classification Thresholds")
body("The five-tier classification uses the following Tavtigian score thresholds:")
thresh_data = [
    ("≥ 10","Pathogenic"),
    ("6 – 9","Likely Pathogenic"),
    ("1 – 5","Variant of Uncertain Significance (VUS)"),
    ("−1 to −6","Likely Benign"),
    ("≤ −7 or BA1","Benign"),
]
th_tbl = doc.add_table(rows=len(thresh_data)+1, cols=2)
th_tbl.style = 'Table Grid'
th_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Total Score","Classification"]):
    th_tbl.rows[0].cells[j].text = h
    for para in th_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, (score, cls) in enumerate(thresh_data):
    th_tbl.rows[i+1].cells[0].text = score
    th_tbl.rows[i+1].cells[1].text = cls
    for cell in th_tbl.rows[i+1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, 11)

heading2("5.5 Flow Diagram")
body(
    "Fig. 5.1 shows the complete data pipeline flowchart from raw ClinVar data ingestion through "
    "preprocessing, feature engineering, model training, and deployment. This flowchart, derived "
    "from the system's research publication, illustrates the end-to-end pipeline that transforms "
    "raw variant records into the trained, calibrated ML model deployed in the production system."
)
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/Screenshot 2026-05-01 at 3.03.10\u202fAM.png",
    "Fig. 5.1: System Data Pipeline Flowchart — ClinVar Download, VUS/UTR Filtering, Stratified Split and SMOTE+Tomek Resampling"
)
body(
    "Fig. 5.2 shows the full ten-layer inference pipeline, tracing a single variant analysis "
    "request from the FastAPI input validation layer through feature engineering, ML inference, "
    "SHAP explainability, ACMG classification, multi-source retrieval, confidence resolution, "
    "contradiction detection, and LLM narrative generation."
)
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/Screenshot 2026-05-01 at 3.03.27\u202fAM.png",
    "Fig. 5.2: Module Chain — Blue Boxes Run on Every Request; Orange Boxes Activate for Pathogenic or Uncertain Predictions"
)
body(
    "Fig. 5.3 presents the simplified ten-layer view of the same pipeline, showing each layer's "
    "role and the data artefact it produces."
)
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/Screenshot 2026-05-01 at 3.04.47\u202fAM.png",
    "Fig. 5.3: Ten-Layer System Architecture — L1 React Frontend to L10 PredictionResult JSON"
)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 6: SYSTEM IMPLEMENTATION
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 6: SYSTEM IMPLEMENTATION")

heading2("6.1 ML Pipeline Implementation")
body(
    "The ML pipeline spans Layers 1–3 of the system architecture and is implemented across "
    "backend/app.py and the data preprocessing scripts."
)

heading2("6.1.1 Dataset")
body(
    "The training dataset was sourced from ClinVar (version 2024), filtered to variants in the "
    "27 epilepsy gene panel. After preprocessing, the dataset comprised approximately 15,000 "
    "labelled variant records with binary labels: Pathogenic/Likely Pathogenic mapped to class 1, "
    "and Benign/Likely Benign mapped to class 0. VUS records were excluded from training to "
    "avoid label noise. The dataset exhibits moderate class imbalance (~60% pathogenic), which "
    "is addressed through stratified k-fold cross-validation and the scale_pos_weight XGBoost parameter."
)

add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/paper/fig7_class_distribution_smote_genes.png",
    "Fig. 6.1: Class Distribution per Gene Before and After SMOTE Balancing"
)

page_break()
heading2("6.1.2 Feature Engineering")
body(
    "The feature engineering pipeline (Layer 2) transforms raw variant attributes into a 93-dimensional "
    "vector. The transformation is deterministic and implemented in the feature_engineer() function "
    "in backend/app.py. Key engineering decisions include:"
)
bullet("Gene encoding: 27 binary flags, one per gene in the panel (SCN1A, SCN2A, SCN3A, SCN8A, SCN9A, KCNQ2, KCNQ3, GABRA1, GABRG2, TSC1, TSC2, MECP2, CDKL5, FOXG1, PCDH19, SLC2A1, SLC6A1, STXBP1, DEPDC5, TBC1D24, LGI1, GRIN2A, CHD2, PRRT2, ALDH7A1, CACNA1A, ARX).")
bullet("Consequence severity: Binary flags for each consequence type plus a severe_consequence_count aggregating frameshift, nonsense, and splice variants.")
bullet("Review score: Ordinal encoding of review status (no criteria=0, single submitter=1, multiple submitters=2, expert panel=3, practice guideline=4).")
bullet("De novo encoding: Binary is_de_novo flag capturing both confirmed and assumed de novo origin.")
bullet("Gene category features: 9 binary flags encoding functional gene families (ion channels, GABA receptors, TSC complex, mTOR pathway, synaptic, etc.) derived from the GENE_MECHANISM map.")
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/paper/17_feature_engineering_pipeline.png",
    "Fig. 6.2: Feature Engineering Pipeline — Raw Input to 93-Dimensional Feature Vector"
)

page_break()
heading2("6.1.3 Model Selection")
body(
    "Six architectures were evaluated on a held-out test fold. Table 6.1 presents the comparative results."
)
p = doc.add_paragraph()
run = p.add_run("Table 6.1: Model Comparison — Accuracy, AUC, F1")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
model_data = [
    ("Logistic Regression", "36.63", "0.5600", "0.5362"),
    ("Random Forest", "87.56", "0.9361", "0.8231"),
    ("Gradient Boosting", "89.64", "0.9503", "0.8465"),
    ("XGBoost (uncalibrated)", "88.90", "0.9468", "0.8464"),
    ("Simple Neural Network", "87.29", "0.9346", "0.8304"),
    ("Residual Neural Network", "87.92", "0.9377", "0.8377"),
    ("XGBoost + Isotonic (Ours)", "89.89", "0.9446", "0.8538"),
]
mdl_tbl = doc.add_table(rows=len(model_data)+1, cols=4)
mdl_tbl.style = 'Table Grid'
mdl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Model", "Accuracy (%)", "AUC", "F1 Score"]):
    mdl_tbl.rows[0].cells[j].text = h
    for para in mdl_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, row_data in enumerate(model_data):
    for j, txt in enumerate(row_data):
        mdl_tbl.rows[i+1].cells[j].text = txt
        for para in mdl_tbl.rows[i+1].cells[j].paragraphs:
            run = para.runs[0] if para.runs else para.add_run(txt)
            set_font(run, 11, bold=(i == len(model_data)-1))

body(
    "Logistic Regression produced degenerate classifier behaviour (accuracy 36.63%) on the tabular "
    "epilepsy feature space, consistent with the non-linear separability of variant features. "
    "Neural network architectures (SimpleNN, ResidualNN) achieved competitive but marginally lower "
    "performance than tree-based ensemble methods, consistent with the literature finding that "
    "gradient-boosted trees outperform deep learning on structured tabular data of this scale. "
    "XGBoost with isotonic calibration achieved the best balance of accuracy (89.89%), well-calibrated "
    "probabilities (Brier score 0.0761 vs. 0.0875 for sigmoid calibration), and interpretability "
    "(SHAP-compatible tree structure)."
)
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/Screenshot 2026-05-01 at 3.04.58\u202fAM.png",
    "Fig. 6.3: Model Training Workflow — SMOTE+Tomek Resampling, XGBoost Fitting, Isotonic Calibration and Val/Test Scoring"
)

heading2("6.1.4 Calibration")
body(
    "Isotonic calibration was applied using scikit-learn's CalibratedClassifierCV with cv=5. "
    "The calibration significantly improved probability reliability: the Brier score improved "
    "from 0.0875 (uncalibrated) to 0.0761 (isotonic). The high-confidence ratio (variants with "
    "predicted probability > 90% or < 10%) is 65.09% of the test set, with accuracy of 96.38% "
    "within that high-confidence stratum — confirming that the model's confident predictions "
    "are highly reliable."
)
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/paper/fig3_ml_validation_test_full.png",
    "Fig. 6.4: ROC Curve (AUC = 0.9446) and Confusion Matrix on Held-Out Test Set"
)
add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/paper/fig2_feature_importance_fixed.png",
    "Fig. 6.5: SHAP Feature Importance — Top Predictive Features for Pathogenicity"
)

page_break()
heading2("6.2 ACMG Engine Implementation")
body(
    "The ACMG engine is implemented in backend/acmg_classifier.py as the ACMGClassifier class. "
    "The classify() method receives gene, consequence, origin, gnomad_data, clinvar_data, "
    "shap_data, variant_type, reference_allele, and alternate_allele, and returns a "
    "comprehensive classification dictionary."
)

page_break()
heading2("6.2.1 Gene Mechanism Map")
body(
    "The GENE_MECHANISM dictionary encodes the primary disease mechanism for each of the 27 panel genes, "
    "distinguishing loss-of-function (LOF) from gain-of-function (GOF) genes. This distinction is "
    "critical for PVS1: truncating variants (frameshift, nonsense, splice) in LOF genes receive PVS1 "
    "(Very Strong, +8 points), but the same variant types in GOF genes (SCN2A, SCN8A, SCN9A) do NOT "
    "receive PVS1 — because in GOF genes, loss-of-function may actually reduce the pathological gain."
)
body("Table 6.2 shows the gene mechanism map for the full 27-gene epilepsy panel.")
p = doc.add_paragraph()
run = p.add_run("Table 6.2: Epilepsy Gene Panel — Disease Mechanism")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
gene_data = [
    ("SCN1A","LOF","Dravet syndrome — haploinsufficiency"),
    ("SCN2A","GOF","Early-onset epilepsy — gain-of-function"),
    ("SCN3A","LOF","Focal epilepsy"),
    ("SCN8A","GOF","EIEE13 — persistent sodium current"),
    ("SCN9A","GOF","Febrile seizures"),
    ("KCNQ2","LOF","Neonatal epilepsy"),
    ("KCNQ3","LOF","Neonatal seizures"),
    ("GABRA1","LOF","Generalised epilepsy"),
    ("GABRG2","LOF","Febrile seizures plus"),
    ("TSC1","LOF","Tuberous sclerosis complex"),
    ("TSC2","LOF","Tuberous sclerosis complex"),
    ("MECP2","LOF","Rett syndrome"),
    ("CDKL5","LOF","Early-onset seizures, X-linked"),
    ("FOXG1","LOF","Rett-like syndrome"),
    ("PCDH19","LOF","Epilepsy limited to females"),
    ("SLC2A1","LOF","GLUT1 deficiency syndrome"),
    ("SLC6A1","LOF","Myoclonic-atonic epilepsy"),
    ("STXBP1","LOF","Ohtahara syndrome"),
    ("DEPDC5","LOF","Focal epilepsy, mTOR pathway"),
    ("TBC1D24","LOF","Multifocal epilepsy"),
    ("LGI1","LOF","Autosomal dominant epilepsy"),
    ("GRIN2A","LOF","Speech-related epilepsy"),
    ("CHD2","LOF","Myoclonic epilepsy"),
    ("PRRT2","LOF","Paroxysmal kinesigenic dyskinesia"),
    ("ALDH7A1","LOF","Pyridoxine-dependent epilepsy"),
    ("CACNA1A","LOF","Absence epilepsy / ataxia"),
    ("ARX","LOF","X-linked epilepsy, infantile spasms"),
]
gene_tbl = doc.add_table(rows=len(gene_data)+1, cols=3)
gene_tbl.style = 'Table Grid'
gene_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Gene","Mechanism","Associated Syndrome"]):
    gene_tbl.rows[0].cells[j].text = h
    for para in gene_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 10, bold=True)
for i, row_data in enumerate(gene_data):
    for j, txt in enumerate(row_data):
        gene_tbl.rows[i+1].cells[j].text = txt
        for para in gene_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 10)

page_break()
heading2("6.2.2 ACMG Validation Results")
body(
    "The ACMG engine was validated against 200 real ClinVar variants with known expert classifications. "
    "Results: overall accuracy 98.5% (197/200 correct); pathogenic accuracy 100% (all 101 pathogenic "
    "variants correctly classified); benign accuracy 97.0% (97/100 benign variants correctly classified). "
    "The tier distribution over the 200 variants: Pathogenic 47 (23.5%), Likely Pathogenic 54 (27.0%), "
    "Likely Benign 97 (48.5%), VUS 2 (1.0%)."
)

page_break()
heading2("6.3 RAG Pipeline Implementation")
body(
    "The Retrieval-Augmented Generation pipeline integrates four data sources to condition LLM generation "
    "on verified clinical evidence, replacing speculative generation with evidence-anchored narrative."
)

heading2("6.3.1 FAISS Literature Index")
body(
    "The PubMed abstract corpus is indexed offline using the PubMedBERT embedding model. Each abstract "
    "is chunked, embedded (768 dimensions), and stored in a FAISS flat L2 index. At runtime, the gene "
    "name (e.g., 'SCN1A epilepsy') is embedded using the same model and the top-5 nearest neighbours "
    "are retrieved. The 363-vector index enables sub-millisecond retrieval."
)

page_break()
heading2("6.3.2 ClinVar and gnomAD Fetchers")
body(
    "The ClinVarFetcher queries NCBI eutils with gene-specific searches and caches results in SQLite. "
    "The GnomADFetcher uses GraphQL to query gnomAD for specific variant identifiers. Both fetchers "
    "implement exponential backoff (2s, 4s, 8s) for API rate limit handling and a 30-day cache TTL "
    "to minimise external API calls during repeated analyses."
)

heading2("6.3.3 Prompt Construction")
body(
    "The RAG generator constructs a multi-section prompt: (1) Variant summary (gene, consequence, "
    "alleles, origin), (2) ML prediction with calibrated probability, (3) ACMG classification with "
    "criteria met and score, (4) SHAP attributions translated to clinical language, (5) gnomAD "
    "population frequency context, (6) ClinVar classification history, (7) Contradiction analysis "
    "if conflicts exist, (8) Retrieved PubMed abstract summaries, (9) PharmGKB drug-gene data. "
    "The LLM is instructed to produce a structured clinical report with sections for Classification, "
    "Evidence Summary, Clinical Significance, and Recommendations."
)

heading2("6.4 User Interface Implementation")
body(
    "The frontend is a React 18 TypeScript single-page application served independently from the "
    "FastAPI backend. The UI is organised into five main components:"
)
bullet("VariantForm: A structured input form with dropdown selectors for gene (27 options), chromosome (1–22, X, Y), molecular consequence (15 options), variant type, review status, and origin. Quick-load buttons populate the form with four pre-defined example variants.")
bullet("ResultCard: Displays the ML prediction (Pathogenic/Benign), calibrated probability as a progress bar, and confidence percentage with colour-coded severity (red for pathogenic, green for benign).")
bullet("ACMGPanel: Presents the ACMG five-tier classification, total Tavtigian score, and a table of all criteria met with their point values and justification text.")
bullet("SHAPPanel: Visualises the top-5 SHAP feature contributors as a ranked list with direction indicators (↑ increases pathogenicity / ↓ decreases) and percentage contributions.")
bullet("ChatInterface: A multi-turn conversational interface where clinicians can ask follow-up questions about the variant, its gene, or the ACMG criteria. The chat system passes the current variant analysis as context to the LLM.")

# ── UI Screenshots ────────────────────────────────────────────────────────────
ui_screenshots = [
    (
        "Screenshot 2026-05-01 at 2.31.12\u202fPM.png",
        "Fig. 6.6: Variant Pathogenicity Analysis — Empty Input Form with SCN1A Gene Selected, Quick Example Buttons and No Prediction Yet Placeholder Panel"
    ),
    (
        "Screenshot 2026-05-01 at 2.32.03\u202fPM.png",
        "Fig. 6.7: SCN1A C→T Nonsense Variant — Filled Input Form with Live Pathogenic Classification (99.5% Confidence), Variant Details and ACMG Score +9 Panel"
    ),
    (
        "Screenshot 2026-05-01 at 2.32.13\u202fPM.png",
        "Fig. 6.8: ACMG/AMP Classification Panel — Evidence Score +9 (Likely Pathogenic), PVS1 Very Strong (+8) and PP3 Supporting (+1) Criteria with LLM Clinical Interpretation"
    ),
    (
        "Screenshot 2026-05-01 at 2.31.27\u202fPM.png",
        "Fig. 6.9: Diagnostic Results Tab — Variant Interpretation (Likely Pathogenic), Contradictory Evidence Analysis and Why This Variant Matters Clinically"
    ),
    (
        "Screenshot 2026-05-01 at 2.31.50\u202fPM.png",
        "Fig. 6.10: RAG-Powered Follow-up Chat — Contraindicated Medications, Additional Considerations and Clinician Query on Sodium Channel Blocker Contraindication with LLM Response"
    ),
]

for idx, (fname, caption) in enumerate(ui_screenshots):
    add_image(
        f"/Users/srinidhisg/epilepsy_diagnostic_assistant/{fname}",
        caption
    )
    # Page break after every 2 images (after index 1, 3) but not after the last
    if idx % 2 == 1 and idx != len(ui_screenshots) - 1:
        page_break()

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 7: TESTING AND RESULTS
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 7: TESTING AND RESULTS")

heading2("7.1 Testing Methodology")
body(
    "A comprehensive testing strategy was designed to validate all aspects of the system prior to "
    "deployment. The strategy follows a bottom-up testing approach, beginning with unit tests on "
    "individual modules, progressing to integration tests that verify inter-module data flow, "
    "and culminating in smoke and performance tests that validate the deployed API."
)
body(
    "All tests are implemented using the pytest 9.0.3 framework and the requests library. "
    "The test suite is organised into five files in the tests/ directory: test_smoke.py, "
    "test_unit.py, test_integration.py, test_regression.py, and test_performance.py. "
    "A conftest.py automatically detects backend availability and skips API-dependent tests "
    "if the backend is not running."
)

heading2("7.2 Test Fixtures")
body("Five clinically representative variant fixtures are used across all test suites:")
fix_data = [
    ("PATHOGENIC_VARIANT","SCN1A","stop_gained","germline","Pathogenic"),
    ("BENIGN_VARIANT","SLC2A1","synonymous_variant","germline","Benign"),
    ("UNCERTAIN_VARIANT","GABRA1","missense_variant","germline","VUS"),
    ("DE_NOVO_VARIANT","KCNQ2","stop_gained","de novo (confirmed)","Pathogenic"),
    ("FRAMESHIFT_VARIANT","TSC2","frameshift_variant","germline","Pathogenic"),
]
fix_tbl = doc.add_table(rows=len(fix_data)+1, cols=5)
fix_tbl.style = 'Table Grid'
fix_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Fixture","Gene","Consequence","Origin","Expected"]):
    fix_tbl.rows[0].cells[j].text = h
    for para in fix_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 10, bold=True)
for i, row_data in enumerate(fix_data):
    for j, txt in enumerate(row_data):
        fix_tbl.rows[i+1].cells[j].text = txt
        for para in fix_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 10)

page_break()
heading2("7.3 Smoke Tests — Results")
body("19 smoke tests were executed. All 19 passed (100% pass rate). Runtime: 6.78 seconds.")
body("Key verifications: /health returns status=ok, model_loaded=True, rag_loaded=True; /predict_variant correctly classifies the PATHOGENIC_VARIANT as Pathogenic and BENIGN_VARIANT as Benign; all endpoints return HTTP 200; invalid input returns HTTP 422.")

page_break()
heading2("7.4 Unit Tests — Results")
body("86 unit tests were executed against the ACMGClassifier, ConfidenceResolver, and ContradictionDetector modules in isolation (no backend required). All 86 passed (100% pass rate). Runtime: 0.42 seconds.")
body("Key verifications: PVS1 fires for stop_gained in all LOF genes; PVS1 does not fire for GOF genes (SCN2A, SCN8A, SCN9A); PS2 fires for confirmed de novo, PM6 for assumed de novo (mutually exclusive); BA1 alone overrides to Benign classification; score arithmetic verified — PVS1=+8, PS2=+4, PM6=+2, BP7=-1; all 6 confidence zones classified correctly; 10 contradiction scenarios detected.")

heading2("7.5 Integration Tests — Results")
body("34 integration tests verified end-to-end data flow between pipeline layers. All 34 passed (100% pass rate). Runtime: 4.05 seconds.")
body("Key verifications: ACMG classification matches ML prediction direction for all 5 fixture variants; SHAP data propagates to ACMG PP3/BP4 criteria; PS2 and PVS1 co-fire correctly for the DE_NOVO_VARIANT; gnomAD frequency criteria (PM2/BA1/BS1) are evaluated in all full-path analyses; RAG explanation field is present and non-empty for pathogenic variants.")

heading2("7.6 Regression Tests — Results")
body("40 regression tests pin known-correct outputs to prevent silent breaking changes. All 40 passed (100% pass rate). Runtime: 0.42 seconds.")
body("Key pins: SCN1A stop_gained → pathogenic_probability > 0.90 (actual: 0.9949); SLC2A1 synonymous → benign_probability > 0.70 (actual: 0.8289); PVS1+PS2+PM2 score = 14 → Pathogenic; all 24 LOF genes verified; all 3 GOF genes verified; UNCERTAIN_LOW=0.3, UNCERTAIN_HIGH=0.7 pinned; all 16 CRITERION_POINTS values pinned.")

page_break()
heading2("7.7 Performance Tests — Results")
body("21 performance tests measured latency, throughput, concurrency, and stability. All 21 passed (100% pass rate). Runtime: 9.32 seconds.")

p = doc.add_paragraph()
run = p.add_run("Table 7.2: Performance SLA Compliance")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
perf_data = [
    ("/health","< 500 ms","2.0 ms","PASS"),
    ("/predict_variant","< 2,000 ms","29–32 ms","PASS"),
    ("/analyze_variant (full RAG)","< 30,000 ms","230–490 ms","PASS"),
    ("/genes","< 500 ms","2.2 ms","PASS"),
    ("/consequences","< 500 ms","1.8 ms","PASS"),
    ("/literature/SCN1A","< 5,000 ms","6–10 ms","PASS"),
]
perf_tbl = doc.add_table(rows=len(perf_data)+1, cols=4)
perf_tbl.style = 'Table Grid'
perf_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Endpoint","SLA Limit","Measured","Status"]):
    perf_tbl.rows[0].cells[j].text = h
    for para in perf_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, row_data in enumerate(perf_data):
    for j, txt in enumerate(row_data):
        perf_tbl.rows[i+1].cells[j].text = txt
        for para in perf_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 11)

body("Throughput: /predict_variant achieves 30–34 req/s sequentially; /health achieves 642 req/s. "
     "Concurrency: 10 concurrent workers achieve 100% success rate (30/30 requests). "
     "Stress: 100 sequential calls produce consistent Pathogenic predictions with p99=28.9 ms. "
     "No response time degradation detected over 50 calls (last-10/first-10 ratio: 0.97×).")

page_break()
heading2("7.8 Overall Test Results Summary")
p = doc.add_paragraph()
run = p.add_run("Table 7.1: Overall Test Results Summary")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
summary_data = [
    ("Smoke Tests","19","19","0","100%","6.78s"),
    ("Unit Tests","86","86","0","100%","0.42s"),
    ("Integration Tests","34","34","0","100%","4.05s"),
    ("Regression Tests","40","40","0","100%","0.42s"),
    ("Performance Tests","21","21","0","100%","9.32s"),
    ("TOTAL","200","200","0","100%","~21s"),
]
sum_tbl = doc.add_table(rows=len(summary_data)+1, cols=6)
sum_tbl.style = 'Table Grid'
sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Suite","Tests","Passed","Failed","Pass Rate","Runtime"]):
    sum_tbl.rows[0].cells[j].text = h
    for para in sum_tbl.rows[0].cells[j].paragraphs:
        for run in para.runs:
            set_font(run, 11, bold=True)
for i, row_data in enumerate(summary_data):
    for j, txt in enumerate(row_data):
        sum_tbl.rows[i+1].cells[j].text = txt
        for para in sum_tbl.rows[i+1].cells[j].paragraphs:
            for run in para.runs:
                set_font(run, 11, bold=(i == len(summary_data)-1))
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 8: CONCLUSION AND FUTURE WORK
# ─────────────────────────────────────────────────────────────────────────────
heading1("CHAPTER 8: CONCLUSION AND FUTURE WORK")

heading2("8.1 Conclusion")
body(
    "This project has successfully designed, implemented, and validated an AI-powered Epilepsy Variant "
    "Diagnostic Assistant that addresses the critical gaps in existing variant classification tools. "
    "The system integrates a calibrated XGBoost machine learning classifier, an automated ACMG/AMP "
    "14-criterion classification engine based on Tavtigian scoring, TreeSHAP-based explainability, "
    "a multi-source RAG pipeline, and a React-based clinical web interface into a coherent, "
    "production-ready diagnostic support system."
)
body(
    "The ML classifier achieves 89.89% accuracy, 0.9446 AUC, and 0.8538 F1 score — outperforming "
    "logistic regression, random forest, gradient boosting, and both simple and residual neural "
    "network architectures on the epilepsy variant classification task. Isotonic calibration "
    "improves the Brier score from 0.0875 to 0.0761, ensuring that model confidence scores are "
    "statistically reliable and clinically meaningful."
)
body(
    "The ACMG engine achieves 98.5% agreement with expert classifications on a validation set of "
    "200 real ClinVar variants, with 100% accuracy on pathogenic variants. The gene mechanism map "
    "correctly differentiates LOF genes (where PVS1 applies) from GOF genes (where PVS1 is "
    "inappropriate), a critical clinical distinction absent from generic variant classifiers."
)
body(
    "Comprehensive testing across 200 test cases (86 unit, 34 integration, 40 regression, 19 smoke, "
    "21 performance) resulted in a 100% pass rate, confirming system correctness, stability, "
    "and production readiness. The API achieves 30–34 req/s throughput for ML inference with "
    "mean latency of 27.6 ms, well within the 2-second SLA requirement."
)
body(
    "The system represents a significant step toward automating the variant interpretation workflow "
    "in epilepsy genetics, with the potential to reduce the manual burden on clinical genetics "
    "laboratories, improve classification consistency, and accelerate diagnosis for patients with "
    "genetic epilepsy."
)

page_break()
heading2("8.2 Future Work")
body(
    "The following directions are identified for future enhancement of the system:"
)
bullet("Whole-Genome Variant Support: Extending the feature engineering pipeline to incorporate genomic position, conservation scores (PhyloP, GERP), and splice prediction scores (SpliceAI) would enable classification of intronic and regulatory variants currently outside the panel's scope.")
bullet("Variant Normalisation: Integration with HGVS notation and Ensembl Variant Effect Predictor (VEP) would enable users to input variants by genomic coordinate or protein change notation rather than requiring manual field entry.")
bullet("Expanded Gene Panel: The current 27-gene panel covers the most clinically established epilepsy genes. Future work should expand to include emerging epilepsy genes identified in recent large-cohort studies (e.g., SYNGAP1, GRIN2B, KCNA2).")
bullet("Longitudinal Variant Tracking: A patient-linked database would enable tracking of variant reclassification over time as new evidence accumulates, automatically notifying clinicians when a VUS is upgraded to Likely Pathogenic.")
bullet("Multi-Tool PP3 Evidence: ACMG PP3 formally requires evidence from multiple independent computational predictors. Future integration of CADD, REVEL, and SpliceAI scores alongside SHAP would enable a more rigorous PP3 assessment.")
bullet("Federated Learning: Training the ML model on multi-institutional variant datasets using federated learning would improve model generalisability while preserving patient data privacy across institutions.")
bullet("Clinician Feedback Loop: A structured feedback mechanism allowing clinical geneticists to flag incorrect classifications would create a continuously improving active-learning loop.")
bullet("PDF Report Generation: Automatic generation of a formatted clinical genetics report in PDF format, suitable for direct inclusion in patient records, is a near-term development priority.")
bullet("Multilingual Support: Extension of the chat interface to support queries in additional languages would broaden accessibility for international clinical users.")
page_break()


# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX A — PROJECT TEAM DETAILS
# ─────────────────────────────────────────────────────────────────────────────
heading1("APPENDIX A: PROJECT TEAM DETAILS")

# Team table: Project Title row + USN/Name/Email/Mobile columns
team_members = [
    ("02JST22UCS110", "Srinidhi S G",      "srinidhisg1@gmail.com",      "8618978068"),
    ("01JST22UCS027", "Bharath H",          "bharath62004@gmail.com",      "9108941433"),
    ("02JST22UCS119", "Sushant Kabbinada",  "sushant.k2150@gmail.com",    "9611130677"),
    ("01JST22UCS008", "Aditya Patil",       "adityapatil0275@gmail.com",  "6363552689"),
]

# Table: 1 project-title row + 1 header row + 4 data rows = 6 rows, 4 cols
app_tbl = doc.add_table(rows=6, cols=4)
app_tbl.style = 'Table Grid'
app_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# ── Row 0: Project Title spanning all 4 columns ──────────────────────────────
title_row = app_tbl.rows[0]
title_row.cells[0].merge(title_row.cells[3])
title_label_cell = title_row.cells[0]
# Left cell: "Project Title" label
# Split: we need 2 cells — label | value. Use a 2-col merge differently:
# Actually re-do as 2-col label+value in first row
# Rebuild: add_table with correct structure
app_tbl._element.getparent().remove(app_tbl._element)

app_tbl = doc.add_table(rows=6, cols=4)
app_tbl.style = 'Table Grid'
app_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths: label | USN-wide | Email-wide | Mobile-narrow
# Total usable width ~5.5 inches (portrait with margins)
col_widths = [Inches(1.3), Inches(1.4), Inches(1.9), Inches(0.9)]
for row in app_tbl.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]

def set_cell_margins(cell, top=100, bottom=100, left=108, right=108):
    """Set cell internal margins in twentieths of a point (twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Row 0: "Project Title" label | merged value (cols 1-3)
r0 = app_tbl.rows[0]
r0.cells[1].merge(r0.cells[3])
cell_pt_label = r0.cells[0]
cell_pt_value = r0.cells[1]

set_cell_margins(cell_pt_label, top=120, bottom=120, left=115, right=115)
set_cell_margins(cell_pt_value, top=120, bottom=120, left=115, right=115)

p_lbl = cell_pt_label.paragraphs[0]
p_lbl.paragraph_format.space_before = Pt(4)
p_lbl.paragraph_format.space_after  = Pt(4)
run_lbl = p_lbl.add_run("Project Title")
set_font(run_lbl, 11, bold=True)
p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT

p_val = cell_pt_value.paragraphs[0]
p_val.paragraph_format.space_before = Pt(4)
p_val.paragraph_format.space_after  = Pt(4)
run_val = p_val.add_run(
    "AI-Powered Epilepsy Variant Diagnostic Assistant Using Machine Learning, "
    "ACMG/AMP Guidelines, and Retrieval-Augmented Generation"
)
set_font(run_val, 11)
p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Row 1: Header row — USN | Team Members | Email ID | Mobile Number
headers = ["USN", "Team Members", "Email ID", "Mobile Number"]
r1 = app_tbl.rows[1]
for j, h in enumerate(headers):
    set_cell_margins(r1.cells[j], top=100, bottom=100, left=115, right=115)
    p = r1.cells[j].paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(h)
    set_font(run, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Rows 2-5: Data rows
for i, (usn, name, email, mobile) in enumerate(team_members):
    r = app_tbl.rows[i + 2]
    for j, txt in enumerate([usn, name, email, mobile]):
        set_cell_margins(r.cells[j], top=100, bottom=100, left=115, right=115)
        p = r.cells[j].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(txt)
        set_font(run, 11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

body("")

# ── Team member photos ────────────────────────────────────────────────────────
# Photos arranged 2 per row with name captions below, like reference image
photo_members = [
    ("1000052397 (1).jpg",                      "Srinidhi S G"),
    ("WhatsApp Image 2026-05-01 at 14.25.35.jpeg", "Bharath H"),
    ("WhatsApp Image 2026-05-01 at 13.58.26.jpeg", "Sushant Kabbinada"),
    ("WhatsApp Image 2026-05-01 at 14.21.19.jpeg", "Aditya Patil"),
]

# 2 photos per row: row0=photo, row1=name, row2=photo, row3=name; 2 cols
photo_tbl = doc.add_table(rows=4, cols=2)
photo_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
photo_tbl.style = 'Normal Table'

for i, (photo_file, name) in enumerate(photo_members):
    row_img  = (i // 2) * 2       # 0 or 2
    row_name = row_img + 1        # 1 or 3
    col      = i % 2              # 0 or 1

    # Photo cell
    photo_cell = photo_tbl.rows[row_img].cells[col]
    pc = photo_cell.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(4)
    if photo_file:
        photo_path = os.path.join(os.path.dirname(__file__), photo_file)
        if os.path.exists(photo_path):
            run_img = pc.add_run()
            run_img.add_picture(photo_path, height=Inches(1.8))

    # Name caption cell
    name_cell = photo_tbl.rows[row_name].cells[col]
    nc = name_cell.paragraphs[0]
    nc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nc.paragraph_format.space_before = Pt(2)
    nc.paragraph_format.space_after  = Pt(12)
    r_name = nc.add_run(name)
    set_font(r_name, 11, bold=True)

body("")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX B — COs, POs AND PSOs MAPPING
# ─────────────────────────────────────────────────────────────────────────────
heading1("APPENDIX B: COs, POs AND PSOs MAPPING FOR THE PROJECT WORK")

heading2("B.1 Course Outcomes (COs)")
body(
    "The following Course Outcomes (COs) are defined for the Project Work (20CS83P). "
    "The justification for how each CO is addressed by this project is provided below."
)

co_defs = [
    ("CO1", "Formulate the problem definition, conduct literature review and apply requirements analysis.",
     "This project addresses CO1 through a rigorous problem formulation identifying the gap in automated epilepsy variant classification. An extensive literature review spanning 20+ peer-reviewed papers on ACMG/AMP guidelines, machine learning for genomics, SHAP explainability, and Retrieval-Augmented Generation was conducted (Chapter 2). A comprehensive requirements analysis producing 12 functional requirements (FR-01 to FR-12) and 10 non-functional requirements (NFR-01 to NFR-10) was completed (Chapter 3)."),
    ("CO2", "Develop and implement algorithms for solving the problem formulated.",
     "CO2 is addressed through the implementation of a ten-layer AI pipeline encompassing: (i) a calibrated XGBoost ML classifier with a 93-dimensional feature engineering pipeline; (ii) an automated ACMG/AMP engine implementing 14 evidence criteria using the Tavtigian Bayesian scoring model; (iii) a TreeSHAP explainability module; (iv) a multi-source RAG pipeline integrating ClinVar, gnomAD, PubMed FAISS indexing, and Groq LLM; and (v) a confidence-aware routing and contradiction detection system. The full implementation is described in Chapters 5 and 6."),
    ("CO3", "Comprehend, present and defend the results of exhaustive testing and explain the major findings.",
     "CO3 is addressed through a comprehensive testing programme comprising 200 test cases across five suites (86 unit, 34 integration, 40 regression, 19 smoke, 21 performance), achieving 100% pass rate. Key results include: ML accuracy 89.89%, AUC 0.9446, F1 0.8538; ACMG accuracy 98.5% on 200 real ClinVar variants; API latency 27.6 ms mean (SLA: 2000 ms); throughput 30–34 req/s. All results are documented in Chapter 7 and defended with statistical evidence (p99 latencies, coefficient of variation, per-gene accuracy analysis)."),
]

for co_id, co_def, justification in co_defs:
    heading3(f"{co_id}: {co_def}")
    body(f"Justification: {justification}")

page_break()
heading2("B.2 Program Outcomes (POs) and Program Specific Outcomes (PSOs)")
body(
    "The following definitions apply to the PO and PSO columns in the mapping table. "
    "The scale used is: 0 = Not Applicable, 1 = Low relevance, 2 = Medium relevance, 3 = High relevance."
)

po_defs = [
    ("PO1","Apply knowledge of computing, mathematics, science, and foundational engineering concepts to solve computer engineering problems."),
    ("PO2","Identify, formulate and analyze complex engineering problems."),
    ("PO3","Plan, implement and evaluate a computer-based system to meet desired societal needs such as economic, environmental, political, healthcare and safety within realistic constraints."),
    ("PO4","Incorporate research methods to design and conduct experiments to investigate real-time problems, to analyze, interpret and provide feasible conclusions."),
    ("PO5","Propose innovative ideas and solutions using modern tools."),
    ("PO6","Apply computing knowledge to assess societal, health, safety, legal and cultural issues and the consequent responsibilities relevant to professional engineering practice."),
    ("PO7","Analyze the local and global impact of computing on individuals and organizations for sustainable development."),
    ("PO8","Adopt ethical principles and uphold the responsibilities and norms of computer engineering practice."),
    ("PO9","Work effectively as an individual and as a member or leader in diverse teams and in multidisciplinary domains."),
    ("PO10","Effectively communicate and comprehend."),
    ("PO11","Demonstrate and apply engineering knowledge and management principles to manage projects in multidisciplinary environments."),
    ("PO12","Recognize contemporary issues and adapt to technological changes for lifelong learning."),
    ("PSO1","Problem Solving Skills: Ability to apply standard practices and mathematical methodologies to solve computational tasks, model real world problems in the areas of database systems, system software, web technologies and Networking solutions with an appropriate knowledge of Data structures and Algorithms."),
    ("PSO2","Knowledge of Computer Systems: An understanding of the structure and working of the computer systems with performance study of various computing architectures."),
    ("PSO3","Successful Career and Entrepreneurship: The ability to get acquainted with the state-of-the-art software technologies leading to entrepreneurship and higher studies."),
    ("PSO4","Computing and Research Ability: Ability to use knowledge in various domains to identify research gaps and to provide solutions to new ideas leading to innovations."),
]
for po_id, po_text in po_defs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run(f"{po_id}: ")
    set_font(r1, 11, bold=True)
    r2 = p.add_run(po_text)
    set_font(r2, 11)

body("")

page_break()
heading2("B.3 Mapping Table: CO, PO and PSO")

p = doc.add_paragraph()
run = p.add_run("Table of Mapping of CO, PO and PSO:")
set_font(run, 12, bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Columns: Subject, Code, CO, PO1..PO12, PSO1..PSO4  = 18 cols total
# Row 0: merged header "Table of Mapping..."  — already done as paragraph above
# Row 1: column headers
# Row 2: Project Work / 20CS83P / CO1 / values
# Row 3: (same subject/code merged) / CO2 / values
# Row 4: (same) / CO3 / values

co_po_data = {
    "CO1": [3,3,2,2,2,2,1,1,2,2,1,2, 3,2,2,2],
    "CO2": [3,3,3,2,3,2,1,1,2,2,2,2, 3,2,2,2],
    "CO3": [2,2,3,2,2,2,2,2,3,3,3,3, 3,2,2,3],
}
po_headers = [f"PO{i}" for i in range(1,13)] + ["PSO1","PSO2","PSO3","PSO4"]

map_tbl = doc.add_table(rows=4, cols=3+len(po_headers))
map_tbl.style = 'Table Grid'
map_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = ["SUBJECT","CODE","CO"] + po_headers
for j, h in enumerate(hdr_cells):
    cell = map_tbl.rows[0].cells[j]
    cell.text = h
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_font(run, 9, bold=True)

# Data rows
for row_idx, (co_name, vals) in enumerate(co_po_data.items()):
    row = map_tbl.rows[row_idx+1]
    if row_idx == 0:
        row.cells[0].text = "Project Work"
        row.cells[1].text = "20CS83P"
    else:
        row.cells[0].text = ""
        row.cells[1].text = ""
    row.cells[2].text = co_name
    for j, v in enumerate(vals):
        row.cells[3+j].text = str(v)
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_font(run, 9)

body("")
body("Note:")
body("Scale: 0 – Not Applicable | 1 – Low relevance | 2 – Medium relevance | 3 – High relevance")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX C — PUBLICATION DETAILS
# ─────────────────────────────────────────────────────────────────────────────
heading1("APPENDIX C: PUBLICATION DETAILS")

add_image(
    "/Users/srinidhisg/epilepsy_diagnostic_assistant/Screenshot 2026-05-01 at 2.24.35\u202fAM.png",
    "",
    width=5.5
)

page_break()




# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
heading1("REFERENCES")

refs = [
    "[1] Richards, S., Aziz, N., Bale, S., Bick, D., Das, S., Gastier-Foster, J., ... & Rehm, H. L. (2015). Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology. Genetics in Medicine, 17(5), 405–424.",
    "[2] Tavtigian, S. V., Greenblatt, M. S., Harrison, S. M., Nussbaum, R. L., Prabhu, S. A., Boucher, K. M., & Biesecker, L. G. (2018). Modeling the ACMG/AMP variant classification guidelines as a Bayesian classification framework. Genetics in Medicine, 20(9), 1054–1060.",
    "[3] Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785–794).",
    "[4] Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765–4774.",
    "[5] Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, P. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459–9474.",
    "[6] Kircher, M., Witten, D. M., Jain, P., O'Roak, B. J., Cooper, G. M., & Shendure, J. (2014). A general framework for estimating the relative pathogenicity of human genetic variants. Nature Genetics, 46(3), 310–315.",
    "[7] Ioannidis, N. M., Rothstein, J. H., Pejaver, V., Middha, S., McDonnell, S. K., Baheti, S., ... & Sadee, W. (2016). REVEL: An ensemble method for predicting the pathogenicity of rare missense variants. American Journal of Human Genetics, 99(4), 877–885.",
    "[8] Karczewski, K. J., Francioli, L. C., Tiao, G., Cummings, B. B., Alföldi, J., Wang, Q., ... & MacArthur, D. G. (2020). The mutational constraint spectrum quantified from variation in 141,456 humans. Nature, 581(7809), 434–443.",
    "[9] Scheffer, I. E., Berkovic, S., Capovilla, G., Connolly, M. B., French, J., Guilhoto, L., ... & Zuberi, S. M. (2017). ILAE classification of the epilepsies: position paper of the ILAE Commission for Classification and Terminology. Epilepsia, 58(4), 512–521.",
    "[10] Landrum, M. J., Lee, J. M., Benson, M., Brown, G. R., Chao, C., Chitipiralla, S., ... & Maglott, D. R. (2018). ClinVar: improving access to variant interpretations and supporting evidence. Nucleic Acids Research, 46(D1), D1062–D1067.",
    "[11] Lundberg, S. M., Erion, G., Chen, H., DeGrave, A., Prutkin, J. M., Nair, B., ... & Lee, S. I. (2020). From local explanations to global understanding with explainable AI for trees. Nature Machine Intelligence, 2(1), 56–67.",
    "[12] Holzinger, A., Biemann, C., Pattichis, C. S., & Kell, D. B. (2017). What do we need to build explainable AI systems for the medical domain? arXiv preprint arXiv:1712.09923.",
    "[13] Nolan, D., Fink, J. K., & de Lonlay, P. (2019). Genomic medicine in epilepsy: Clinical impact and diagnostic challenges. The Lancet Neurology, 18(5), 504–515.",
    "[14] Xiong, G., Jin, Q., Lu, Z., & Zhang, A. (2023). Benchmarking retrieval-augmented generation for medicine. arXiv preprint arXiv:2402.13178.",
    "[15] Stenson, P. D., Ball, E. V., Mort, M., Phillips, A. D., Shiel, J. A., Thomas, N. S., ... & Cooper, D. N. (2003). Human Gene Mutation Database (HGMD): 2003 update. Human Mutation, 21(6), 577–581.",
    "[16] McLaren, W., Gil, L., Hunt, S. E., Riat, H. S., Ritchie, G. R., Thormann, A., ... & Cunningham, F. (2016). The ensembl variant effect predictor. Genome Biology, 17(1), 122.",
    "[17] Adzhubei, I. A., Schmidt, S., Peshkin, L., Ramensky, V. E., Gerasimova, A., Bork, P., ... & Sunyaev, S. R. (2010). A method and server for predicting damaging missense mutations. Nature Methods, 7(4), 248–249.",
    "[18] Whiffin, N., Minikel, E., Walsh, R., O'Donnell-Luria, A. H., Karczewski, K., Ing, A. Y., ... & Watkins, H. (2017). Using high-resolution variant frequencies to empower clinical genome interpretation. Genetics in Medicine, 19(10), 1151–1158.",
    "[19] Lek, M., Karczewski, K. J., Minikel, E. V., Samocha, K. E., Banks, E., Fennell, T., ... & MacArthur, D. G. (2016). Analysis of protein-coding genetic variation in 60,706 humans. Nature, 536(7616), 285–291.",
    "[20] Johnson, A. E., Pollard, T. J., Shen, L., Lehman, L. W. H., Feng, M., Ghassemi, M., ... & Mark, R. G. (2023). Large language models in medical and clinical informatics: State of the art and future directions. Journal of Medical Internet Research, 25, e47200.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(ref)
    set_font(run, 11)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
output_path = "/Users/srinidhisg/epilepsy_diagnostic_assistant/Epilepsy_Diagnostic_Assistant_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
